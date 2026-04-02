# Fresh trigger for Railway Deployment - v11.0.1
import io
# ระบบถูกปรับมาใช้ Gunicorn gthread worker เพื่อความเสถียรสูงสุดบน Python 3.12+
# และแก้ไขปัญหา [CRITICAL] WORKER TIMEOUT โดยการใช้ Threading มาตรฐานครับ
print("🚀 SYSTEM v11.0 [GTHREAD/STREAM] READY.", flush=True)

# Force UTF-8 output encoding for Windows compatibility
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
if sys.stderr.encoding != 'utf-8':
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

from dotenv import load_dotenv
from pathlib import Path
BASE_DIR = Path(__file__).parent.absolute()
load_dotenv(BASE_DIR / ".env", override=True)

import os
import uuid
from pathlib import Path

from flask import Flask, request, jsonify, render_template, render_template_string, send_from_directory, session, send_file
import hmac
import hashlib
import base64
from functools import wraps
from flask_cors import CORS
from flask_socketio import SocketIO, emit, join_room, leave_room
from werkzeug.utils import secure_filename
from bs4 import BeautifulSoup
import requests

import json
import time
from datetime import datetime
import sqlite3
import threading
# ใช้ google-genai (SDK ใหม่) แทน google-generativeai (deprecated)
from google import genai
import rag_engine
import database
import export_service
import ai_providers
import notification_db
from fpdf import FPDF
from pywebpush import webpush, WebPushException

try:
    from google.oauth2 import id_token
    from google.auth.transport import requests as google_requests
    GOOGLE_AUTH_AVAILABLE = True
except ImportError:
    GOOGLE_AUTH_AVAILABLE = False


# --- VAPID KEYS FOR PUSH NOTIFICATIONS ---
VAPID_PUBLIC_KEY = os.environ.get("VAPID_PUBLIC_KEY", "BNW7f7p3Ush_rg9vjIXxz1KTthTsiy3rz17oaygTy1-l4bTQJKpLeYEj4v3jYQkggo1VLa7w7sNb6mWDaIVH5eU")
VAPID_PRIVATE_KEY = os.environ.get("VAPID_PRIVATE_KEY", "iVj1ybjnV5yP34d1BZ9ydP6Ad_m_24FC6AhYkc2On04")
VAPID_CLAIMS = {"sub": "mailto:fewit@example.com"}

from concurrent.futures import ThreadPoolExecutor

def send_push_notification(username, title, message, url='/', tag=None):
    """Send a real-time push notification to a user's device via VAPID."""
    subscriptions = database.get_push_subscriptions(username)
    if not subscriptions:
        return

    payload_data = {
        "title": title,
        "body": message,
        "url": url
    }
    if tag:
        payload_data["tag"] = tag
        
    payload = json.dumps(payload_data)

    for sub_json in subscriptions:
        try:
            subscription_info = json.loads(sub_json)
            webpush(
                subscription_info=subscription_info,
                data=payload,
                vapid_private_key=VAPID_PRIVATE_KEY,
                vapid_claims=VAPID_CLAIMS.copy()
            )
        except WebPushException as ex:
            print(f"WebPush error for {username}: {ex}")
            if ex.response is not None and ex.response.status_code in [404, 410]:
                database.remove_push_subscription(username, sub_json)
        except Exception as e:
            # If it's the curve error, we might need to log more details but don't crash the worker
            print(f"General Push error for {username}: {e}")

def batch_send_push_notification(usernames, title, message, url='/'):
    """Send push notifications to multiple users in parallel."""
    if not usernames:
        return
    with ThreadPoolExecutor(max_workers=10) as executor:
        for uname in usernames:
            executor.submit(send_push_notification, uname, title, message, url)

# Re-init KB to catch settings after load_dotenv
rag_engine.reinit_kb()

database.init_db()

# Hardcoded credentials for internal use
USERS = {
    "admin": "1234",
    "few": "few1234",
    "do": "do1234"
}

# Initialize "General" group if it doesn't exist
conn = sqlite3.connect(database.DB_PATH)
cursor = conn.cursor()
cursor.execute("SELECT id FROM chat_rooms WHERE id = 1")
if not cursor.fetchone():
    cursor.execute("INSERT INTO chat_rooms (id, name, owner) VALUES (1, 'กลุ่มทั่วไป (General)', 'System')")
    
# Always ensure hardcoded users and db users are in General room (ID: 1)
all_usernames = set([u.capitalize() for u in USERS.keys()])
all_usernames.add("AI-Assistant") # Bot member
for u in database.get_all_usernames():
    all_usernames.add(u)

# Ensure AI Assistant profile exists
cursor.execute("INSERT OR IGNORE INTO user_profiles (username, display_name, avatar_url) VALUES ('AI-Assistant', 'AI Assistant', 'https://cdn-icons-png.flaticon.com/512/4712/4712035.png')")

# Ensure Admin has admin role
cursor.execute("INSERT OR IGNORE INTO user_settings (username, role) VALUES ('Admin', 'admin')")
cursor.execute("UPDATE user_settings SET role = 'admin' WHERE username = 'Admin'")

for u in all_usernames:
    cursor.execute("INSERT OR IGNORE INTO room_members (room_id, username) VALUES (1, ?)", (u,))

conn.commit()
conn.close()

# --- Thai Holidays 2026 ---
THAI_HOLIDAYS_2026 = {
    "2026-01-01": "วันขึ้นปีใหม่",
    "2026-02-11": "วันมาฆบูชา",
    "2026-04-06": "วันจักรี",
    "2026-04-13": "วันสงกรานต์",
    "2026-04-14": "วันสงกรานต์",
    "2026-04-15": "วันสงกรานต์",
    "2026-05-01": "วันแรงงานแห่งชาติ",
    "2026-05-04": "วันฉัตรมงคล",
    "2026-05-11": "วันพืชมงคล",
    "2026-05-20": "วันวิสาขบูชา",
    "2026-06-03": "วันเฉลิมพระชนมพรรษาพระราชินี",
    "2026-07-18": "วันอาสาฬหบูชา",
    "2026-07-19": "วันเข้าพรรษา",
    "2026-07-28": "วันเฉลิมพระชนมพรรษา ร.10",
    "2026-08-12": "วันแม่แห่งชาติ",
    "2026-10-13": "วันนวมินทรมหาราช",
    "2026-10-23": "วันปิยมหาราช",
    "2026-12-05": "วันคล้ายวันเฉลิมพระชนมพรรษา ร.9 / วันพ่อแห่งชาติ",
    "2026-12-10": "วันรัฐธรรมนูญ",
    "2026-12-31": "วันสิ้นปี"
}

def get_current_time():
    import datetime
    now = datetime.datetime.now()
    return now.strftime("%Y-%m-%d %H:%M:%S")

# ─── Weather Context for AI (Background Updated) ───────────
_weather_cache = {"data": "กำลังเตรียมข้อมูลสภาพอากาศ...", "timestamp": 0}

def update_weather_background():
    """Fetch Bangkok weather in background to prevent blocking chat."""
    global _weather_cache
    import time
    import requests
    
    WMO_CODES = {
        0: "ท้องฟ้าแจ่มใส", 1: "ท้องฟ้าโปร่ง", 2: "มีเมฆบางส่วน", 3: "เมฆมาก",
        45: "หมอกลง", 48: "หมอกน้ำค้างแข็ง",
        51: "ฝนปรอยๆ", 53: "ฝนปรอยๆ", 55: "ฝนปรอยๆ หนาแน่น",
        61: "ฝนตกเล็กน้อย", 63: "ฝนตกปานกลาง", 65: "ฝนตกหนัก",
        80: "ฝนซู่เล็กน้อย", 81: "ฝนซู่ปานกลาง", 82: "ฝนซู่รุนแรง",
        95: "พายุฝนฟ้าคะนอง",
    }
    
    try:
        t_start = time.time()
        w_url = "https://api.open-meteo.com/v1/forecast?latitude=13.75&longitude=100.51&current=temperature_2m,weather_code&timezone=Asia%2FBangkok"
        aq_url = "https://air-quality-api.open-meteo.com/v1/air-quality?latitude=13.75&longitude=100.51&current=pm2_5&timezone=Asia%2FBangkok"
        
        w_res = requests.get(w_url, timeout=10).json()
        aq_res = requests.get(aq_url, timeout=10).json()
        
        cur = w_res.get("current", {})
        temp = round(cur.get("temperature_2m", 0))
        cond = WMO_CODES.get(cur.get("weather_code", 0), "แจ่มใส")
        pm25 = aq_res.get("current", {}).get("pm2_5", 0)
        
        _weather_cache["data"] = f"ขณะนี้ในกรุงเทพฯ: {cond}, อุณหภูมิ {temp}°C, PM2.5: {pm25} µg/m³"
        _weather_cache["timestamp"] = time.time()
        print(f"🌤️ Background Weather Synced ({time.time() - t_start:.2f}s)", flush=True)
    except Exception as e:
        print(f"⚠️ Background Weather Error: {e}", flush=True)

def get_weather_context():
    """Return cached weather immediately from background sync."""
    return _weather_cache["data"]


app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "orgchat-super-secret-key-1234")
CORS(app, supports_credentials=True)
# บังคับ async_mode เป็น threading เพื่อความเสถียรและทำงานร่วมกับ gthread ได้ดีที่สุด
socketio = SocketIO(app, cors_allowed_origins="*", async_mode='threading', max_http_payload_size=10*1024*1024)

VERSION = "1.10.0-STABLE"

# ─── APScheduler: Daily Summary at 08:00 ───────────────────
try:
    from apscheduler.schedulers.background import BackgroundScheduler
    _scheduler = BackgroundScheduler(daemon=True)
    _scheduler.add_job(update_weather_background, 'interval', minutes=30)
    try:
        _scheduler.start()
        # Run once at boot (softly)
        update_weather_background()
    except:
        pass

    def _scheduled_daily_summary():
        """Run daily AI summary and push to all users at 08:00."""
        try:
            data = database.get_daily_activities()
            posts = data.get("posts", [])
            schedules = data.get("schedules", [])
            if not posts and not schedules:
                return
            posts_text = "\n".join([f"- {p['author']} โพสต์: {p['content'][:80]}" for p in posts])
            sched_text = "\n".join([f"- {s['title']} วันที่ {s['date']} {s['time']}" for s in schedules])
            prompt = (
                "สรุปกิจกรรมสำคัญประจำวันในบริษัทให้เพื่อนร่วมงานฟังแบบเป็นกันเองแต่เป็นมืออาชีพ (ไม่เกิน 3 ประโยค) ใช้สำนวนภาษาไทยที่เป็นธรรมชาติเหมือนเพื่อนร่วมงานเล่าให้กันฟัง:\n"
                f"โพสต์:{posts_text}\nตาราง:{sched_text}"
            )
            provider = ai_providers.get_provider()
            summary = ""
            for chunk in provider.chat_stream(prompt, [], "คุณคือ OrgChat AI ผู้ช่วยสรุปข่าวสารองค์กร"):
                if chunk:
                    summary += chunk
            if summary:
                unames = database.get_all_usernames()
                notification_db.notify_users(
                    unames, "daily_summary", "📋 Daily Briefing",
                    summary[:200], link="#feed"
                )
                batch_send_push_notification(unames, "📋 Daily Briefing", summary[:120], url="/")
                
                # --- Send to LINE as well ---
                threading.Thread(target=broadcast_line_announcement, args=(
                    "สรุปข่าวสารองค์กร", 
                    summary[:400]
                ), kwargs={
                    "fields": {
                        "วันที่": datetime.datetime.now().strftime("%Y-%m-%d"),
                        "สรุป": summary[:300] + ("..." if len(summary) > 300 else "")
                    }
                }).start()
                
                print(f"✅ Daily summary sent to {len(unames)} users")
        except Exception as e:
            print(f"❌ Scheduled daily summary error: {e}")

    _scheduler.add_job(_scheduled_daily_summary, 'cron', hour=8, minute=0, id='daily_summary')
except ImportError:
    print("⚠️ APScheduler not installed — run: pip install APScheduler")

@app.route("/api/debug/routes")
def list_routes():
    import urllib.parse
    output = []
    for rule in app.url_map.iter_rules():
        methods = ','.join(rule.methods)
        line = urllib.parse.unquote(f"{rule.endpoint:25} {methods:20} {rule}")
        output.append(line)
    return "<pre>" + "\n".join(sorted(output)) + "</pre>"

# ─── SocketIO Online Status Tracking ───────────────
online_users_registry = {} # sid -> username

@socketio.on('connect')
def handle_connect():
    print(f"Client connected: {request.sid}")

@socketio.on('disconnect')
def handle_disconnect():
    if request.sid in online_users_registry:
        user = online_users_registry.pop(request.sid)
        print(f"User disconnected: {user}")
        # Broadcast updated list of unique online usernames
        active_users = list(set(online_users_registry.values()))
        socketio.emit('user_status_updated', {'online_users': active_users}, broadcast=True)

# ─── SocketIO Event Handlers ───────────────────────
@socketio.on('join')
def on_join(data):
    room = data.get('room')
    if room:
        join_room(room)

@socketio.on('leave')
def on_leave(data):
    room = data.get('room')
    if room:
        leave_room(room)

@socketio.on('user_online')
def on_user_online(data):
    # This event is emitted by the client strictly for tracking online status
    username = data.get('username')
    if username:
        online_users_registry[request.sid] = username
        join_room(f"user_{username}")
        # Broadcast all unique online users
        active_users = list(set(online_users_registry.values()))
        socketio.emit('user_status_updated', {'online_users': active_users}, broadcast=True)

@socketio.on('user_typing')
def on_user_typing(data):
    room_id = data.get('roomName')
    if room_id:
        # Assuming we just broadcast it, the UI handles filtering 
        emit('user_typing', data, broadcast=True)



@app.route("/api/version")
def get_version():
    return jsonify({"version": VERSION, "status": "online"})

ALLOWED_EXTENSIONS = {".pdf", ".csv", ".txt", ".md", ".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp", ".docx", ".xlsx"}


def _get_gemini_api_key() -> str | None:
    return os.environ.get("GEMINI_API_KEY", "").strip() or None


def _configure_gemini(api_key: str):
    # google-genai ใหม่สร้าง client ต่อ session — ไม่มี configure() แบบ global
    pass  # ไม่จำเป็นต้อง configure global แล้ว


def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if "user" not in session:
            return jsonify({"ok": False, "error": "กรุณาเข้าสู่ระบบก่อนใช้งาน"}), 401
        return f(*args, **kwargs)
    return decorated_function


def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if "user" not in session:
            return jsonify({"ok": False, "error": "กรุณาเข้าสู่ระบบก่อนใช้งาน"}), 401
        
        user = session.get("user")
        settings = database.get_user_setting(user)
        if settings.get("role") != "admin":
            return jsonify({"ok": False, "error": "เฉพาะผู้ดูแลระบบเท่านั้นที่มีสิทธิ์เข้าถึงส่วนนี้"}), 403
        return f(*args, **kwargs)
    return decorated_function

def can_edit_knowledge_base():
    user = session.get("user")
    if not user: return False
    settings = database.get_user_setting(user)
    if settings.get("role") == "admin":
        return True
    return bool(settings.get("can_edit_kb", False))

def can_view_knowledge_base():
    user = session.get("user")
    if not user: return False
    settings = database.get_user_setting(user)
    if settings.get("role") == "admin":
        return True
    return bool(settings.get("can_view_kb", False))

def can_delete_knowledge_base():
    user = session.get("user")
    if not user: return False
    settings = database.get_user_setting(user)
    if settings.get("role") == "admin":
        return True
    return bool(settings.get("can_delete_kb", False))

def is_admin():
    user = session.get("user")
    if not user: return False
    settings = database.get_user_setting(user)
    return settings.get("role") == "admin"

def get_rag_filter(username):
    """Constructs a ChromaDB filter to restrict AI search based on category permissions."""
    if not username:
        return None # Fallback
        
    # Admin bypass
    user_settings = database.get_user_setting(username)
    if user_settings.get("role") == "admin":
        return None
        
    # Get categories the user can access
    cats = database.get_categories(username)
    # Store IDs as both int and str to bridge any type mismatch in ChromaDB
    allowed_ids = []
    for c in cats:
        if c.get('id') is not None:
            allowed_ids.append(c['id'])
            allowed_ids.append(str(c['id']))
    
    # We allow unassigned files (empty string in ChromaDB) to all users who have KB access
    unassigned_filters = [{"category_id": ""}]
    
    if not allowed_ids:
        # User has no specific category access, only unassigned
        return {"$or": unassigned_filters}
    
    # Return filter: category_id in allowed_ids OR category_id is unassigned
    res = {
        "$or": [
            {"category_id": {"$in": allowed_ids}},
            *unassigned_filters
        ]
    }
    # Note: log_bot is defined later in the file, so we check if it exists or call it safely
    if 'log_bot' in globals() or 'log_bot' in locals():
        log_bot(f"🛡️ [AUTH] RAG filter for {username}: {len(allowed_ids)//2} allowed categories.")
    else:
        print(f"🛡️ [AUTH] RAG filter for {username}: {len(allowed_ids)//2} allowed categories.")
    return res


# ─────────────── Routes ───────────────

@app.route('/uploads/social_feed/<path:filename>')
def serve_social_uploads(filename):
    return send_from_directory('uploads/social_feed', filename)

@app.route('/uploads/profiles/<path:filename>')
def serve_profile_uploads(filename):
    return send_from_directory('uploads/profiles', filename)

@app.route('/uploads/group_chat/<path:filename>')
def serve_group_chat_uploads(filename):
    return send_from_directory('uploads/group_chat', filename)

@app.route('/uploads/dm_chat/<path:filename>')
def serve_dm_chat_uploads(filename):
    return send_from_directory('uploads/dm_chat', filename)

@app.route('/uploads/group_profiles/<path:filename>')
def serve_group_profile_uploads(filename):
    return send_from_directory('uploads/group_profiles', filename)

@app.route("/api/link-preview")
@login_required
def get_link_preview():
    url = request.args.get("url")
    if not url:
        return jsonify({"ok": False}), 400
    if not url.startswith("http"):
        url = "http://" + url
        
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
        response = requests.get(url, timeout=5, headers=headers)
        soup = BeautifulSoup(response.text, 'html.parser')
        
        title = soup.find("title").text if soup.find("title") else url
        description = ""
        og_desc = soup.find("meta", property="og:description")
        if og_desc:
            description = og_desc["content"]
        else:
            meta_desc = soup.find("meta", attrs={"name": "description"})
            if meta_desc:
                description = meta_desc["content"]
                
        image = ""
        og_image = soup.find("meta", property="og:image")
        if og_image:
            image = og_image["content"]
            
        return jsonify({
            "ok": True,
            "title": title[:100],
            "description": description[:200],
            "image": image,
            "url": url
        })
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/login", methods=["POST"])
def api_login():
    data = request.json
    username_raw = data.get("username", "").strip()
    username_lower = username_raw.lower()
    password = data.get("password", "")
    
    # 1. Check Database first (Dynamic Users) — supports bcrypt hashed passwords
    settings = database.get_user_setting(username_raw)
    
    if settings.get("custom_password") and database.check_password(password, settings["custom_password"]):
        if not settings.get("is_active", 1):
            return jsonify({"ok": False, "error": "บัญชีนี้ถูกระงับโดยผู้ดูแลระบบ"}), 403
        session.permanent = True
        session["user"] = settings.get("username_original", username_raw)
        return jsonify({"ok": True, "user": session["user"]})
        
    # 2. Fallback to Hardcoded (plain text)
    if username_lower in USERS and USERS[username_lower] == password:
        user_to_session = username_raw.capitalize()
        if not settings.get("is_active", 1):
             return jsonify({"ok": False, "error": "บัญชีนี้ถูกระงับโดยผู้ดูแลระบบ"}), 403
        session.permanent = True
        session["user"] = user_to_session
        return jsonify({"ok": True, "user": user_to_session})

    return jsonify({"ok": False, "error": "ชื่อผู้ใช้หรือรหัสผ่านไม่ถูกต้อง"}), 401

# ─── QR Code Login System ─────────────────────────────────
import uuid
_qr_tokens = {}  # token -> {"status": "pending"|"approved", "user": None, "created": timestamp}

@app.route("/api/qr/generate", methods=["POST"])
def qr_generate():
    """Generate a unique QR login token for desktop."""
    token = str(uuid.uuid4())
    _qr_tokens[token] = {
        "status": "pending",
        "user": None,
        "created": time.time()
    }
    # Cleanup old tokens (> 5 min)
    cutoff = time.time() - 300
    expired = [k for k, v in _qr_tokens.items() if v["created"] < cutoff]
    for k in expired:
        del _qr_tokens[k]
    
    return jsonify({"ok": True, "token": token})

@app.route("/api/qr/poll/<token>")
def qr_poll(token):
    """Desktop polls this to check if QR was scanned and approved."""
    entry = _qr_tokens.get(token)
    if not entry:
        return jsonify({"ok": False, "error": "Token expired or invalid"}), 404
    
    # Check expiry (5 min)
    if time.time() - entry["created"] > 300:
        del _qr_tokens[token]
        return jsonify({"ok": False, "error": "Token expired"}), 410
    
    if entry["status"] == "approved" and entry["user"]:
        user = entry["user"]
        del _qr_tokens[token]
        # Log user in
        session.permanent = True
        session["user"] = user
        return jsonify({"ok": True, "status": "approved", "user": user})
    
    return jsonify({"ok": True, "status": "pending"})

@app.route("/api/qr/approve/<token>", methods=["POST"])
@login_required
def qr_approve(token):
    """Mobile user (already logged in) approves a QR login token."""
    entry = _qr_tokens.get(token)
    if not entry:
        return jsonify({"ok": False, "error": "Token ไม่ถูกต้องหรือหมดอายุ"}), 404
    
    if time.time() - entry["created"] > 300:
        del _qr_tokens[token]
        return jsonify({"ok": False, "error": "Token หมดอายุแล้ว"}), 410
    
    user = session.get("user")
    entry["status"] = "approved"
    entry["user"] = user
    return jsonify({"ok": True, "message": f"อนุมัติเข้าสู่ระบบสำหรับ {user} เรียบร้อย"})

@app.route("/qr-login/<token>")
def qr_login_page(token):
    """Page that mobile opens after scanning QR code."""
    entry = _qr_tokens.get(token)
    if not entry:
        return render_template_string(QR_APPROVE_HTML, token=token, error="Token ไม่ถูกต้องหรือหมดอายุ", valid=False)
    if time.time() - entry["created"] > 300:
        return render_template_string(QR_APPROVE_HTML, token=token, error="Token หมดอายุแล้ว กรุณาสร้าง QR Code ใหม่", valid=False)
    user = session.get("user")
    if not user:
        return render_template_string(QR_APPROVE_HTML, token=token, error="กรุณาเข้าสู่ระบบบนมือถือก่อน แล้วสแกนอีกครั้ง", valid=False)
    return render_template_string(QR_APPROVE_HTML, token=token, error=None, valid=True, user=user)

QR_APPROVE_HTML = """<!DOCTYPE html>
<html lang="th">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>OrgChat QR Login</title>
<style>
  * { margin: 0; padding: 0; box-sizing: border-box; }
  body { font-family: 'Inter', -apple-system, sans-serif; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); min-height: 100vh; display: flex; align-items: center; justify-content: center; padding: 20px; }
  .card { background: white; border-radius: 2rem; padding: 3rem 2rem; max-width: 400px; width: 100%; text-align: center; box-shadow: 0 25px 80px rgba(0,0,0,0.3); }
  .icon { width:80px; height:80px; border-radius:50%; display:flex; align-items:center; justify-content:center; margin:0 auto 1.5rem; font-size:2.5rem; }
  .icon.success { background: #ecfdf5; }
  .icon.error { background: #fef2f2; }
  h1 { font-size:1.3rem; font-weight:800; color:#1e293b; margin-bottom:.5rem; }
  p { font-size:.85rem; color:#64748b; margin-bottom:1.5rem; line-height:1.6; }
  .user-badge { background:linear-gradient(135deg,#667eea,#764ba2); color:white; padding:.5rem 1.5rem; border-radius:999px; font-weight:700; font-size:.9rem; display:inline-block; margin-bottom:1.5rem; }
  .btn { display:block; width:100%; padding:1rem; border-radius:1rem; border:none; font-size:1rem; font-weight:800; cursor:pointer; transition:all .2s; }
  .btn-primary { background:linear-gradient(135deg,#667eea,#764ba2); color:white; }
  .btn-primary:hover { transform:scale(1.02); box-shadow:0 8px 25px rgba(102,126,234,.4); }
  .btn-primary:disabled { opacity:.6; cursor:not-allowed; transform:none; }
  .btn-secondary { background:#f1f5f9; color:#475569; margin-top:.8rem; }
  .done { display:none; }
  .done.show { display:block; }
  .done .icon { background:#ecfdf5; }
</style>
</head>
<body>
<div class="card">
  {% if not valid %}
    <div class="icon error">❌</div>
    <h1>ไม่สามารถดำเนินการได้</h1>
    <p>{{ error }}</p>
    <a href="/" class="btn btn-secondary" style="text-decoration:none;">กลับหน้าหลัก</a>
  {% else %}
    <div id="approveView">
      <div class="icon success">📱</div>
      <h1>ยืนยันการเข้าสู่ระบบ</h1>
      <p>อุปกรณ์อื่นกำลังขอเข้าสู่ระบบด้วยบัญชีของคุณ</p>
      <div class="user-badge">👤 {{ user }}</div>
      <br><br>
      <button id="approveBtn" class="btn btn-primary" onclick="approveLogin()">✅ อนุมัติเข้าสู่ระบบ</button>
      <button class="btn btn-secondary" onclick="window.close()">ยกเลิก</button>
    </div>
    <div id="doneView" class="done">
      <div class="icon success">✅</div>
      <h1>อนุมัติเรียบร้อยแล้ว!</h1>
      <p>อุปกรณ์อีกเครื่องจะเข้าสู่ระบบอัตโนมัติ<br>คุณสามารถปิดหน้านี้ได้</p>
    </div>
  {% endif %}
</div>
<script>
async function approveLogin() {
  const btn = document.getElementById('approveBtn');
  btn.disabled = true;
  btn.textContent = 'กำลังอนุมัติ...';
  try {
    const res = await fetch('/api/qr/approve/{{ token }}', { method: 'POST' });
    const data = await res.json();
    if (data.ok) {
      document.getElementById('approveView').style.display = 'none';
      document.getElementById('doneView').classList.add('show');
    } else {
      alert(data.error || 'เกิดข้อผิดพลาด');
      btn.disabled = false;
      btn.textContent = '✅ อนุมัติเข้าสู่ระบบ';
    }
  } catch(e) {
    alert('เกิดข้อผิดพลาดในการเชื่อมต่อ');
    btn.disabled = false;
    btn.textContent = '✅ อนุมัติเข้าสู่ระบบ';
  }
}
</script>
</body>
</html>"""

@app.route("/api/login/google", methods=["POST"])
def api_login_google():
    data = request.json
    token = data.get("credential")
    client_id = os.environ.get("GOOGLE_CLIENT_ID", "")
    
    if not token or not client_id:
        return jsonify({"ok": False, "error": "ไม่สามารถตรวจสอบสิทธิ์ด้วย Google ได้ (ไม่มี Token หรือ Client ID)"}), 400

    if not GOOGLE_AUTH_AVAILABLE:
        return jsonify({"ok": False, "error": "Google Auth library not installed on server"}), 500

    try:
        # Verify the token
        idinfo = id_token.verify_oauth2_token(token, google_requests.Request(), client_id)

        # Token is valid, get the user's email
        email = idinfo.get("email")
        if not email:
            return jsonify({"ok": False, "error": "ไม่พบอีเมลจาก Google Account"}), 400
            
        username = email.split('@')[0]
        
        # Check if user exists or is active
        settings = database.get_user_setting(username)
        if settings and not settings.get("is_active", 1):
             return jsonify({"ok": False, "error": "บัญชีนี้ถูกระงับโดยผู้ดูแลระบบ"}), 403

        # Update user profile with Google info if they have no avatar/display name
        profile = database.get_user_profile(username)
        display_name = profile.get("display_name") if profile.get("display_name") else idinfo.get("name")
        avatar_url = profile.get("avatar_url") if profile.get("avatar_url") else idinfo.get("picture")
        
        database.update_user_profile(username, display_name=display_name, avatar_url=avatar_url)

        session.permanent = True
        session["user"] = username
        return jsonify({"ok": True, "user": username})

    except Exception as e:
        print(f"Google Login Error: {e}")
        return jsonify({"ok": False, "error": "การตรวจสอบสิทธิ์กับ Google ล้มเหลว"}), 401

# --- LINE BOT HELPERS ---
def verify_line_signature(body: bytes, signature: str) -> bool:
    """
    ตรวจสอบ signature จาก LINE Messaging API
    ตาม spec: HMAC-SHA256(channel_secret, body) → Base64
    """
    channel_secret = os.environ.get("LINE_CHANNEL_SECRET", "").strip()
    if not channel_secret:
        print("❌ [LINE] ไม่พบ LINE_CHANNEL_SECRET ใน Environment Variables!", flush=True)
        return False

    # คำนวณ HMAC-SHA256 ตาม LINE API spec
    hash_val = hmac.new(
        channel_secret.encode('utf-8'),
        body,  # ต้องเป็น raw bytes ก่อน request.json ถูกอ่าน
        hashlib.sha256
    ).digest()
    expected_signature = base64.b64encode(hash_val).decode('utf-8')

    # clean signature จาก header (กำจัด whitespace/newline ที่ LINE อาจส่งมา)
    received_signature = signature.strip()

    # เปรียบเทียบแบบ constant-time เพื่อความปลอดภัย
    try:
        match = hmac.compare_digest(received_signature, expected_signature)
    except (TypeError, ValueError):
        match = False

    if not match:
        print(f"🛡️ [LINE] Signature ไม่ตรง!", flush=True)
        print(f"  → รับมา  : {received_signature[:20]}...", flush=True)
        print(f"  → คาดหวัง: {expected_signature[:20]}...", flush=True)
        print(f"  → SECRET ยาว {len(channel_secret)} ตัวอักษร", flush=True)

    return match

def reply_to_line(reply_token, text):
    token = os.environ.get("LINE_CHANNEL_ACCESS_TOKEN", "").strip()
    if not token or not reply_token: return
    
    url = "https://api.line.me/v2/bot/message/reply"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {token}"
    }
    payload = {
        "replyToken": reply_token,
        "messages": [{"type": "text", "text": text}]
    }
    try:
        res = requests.post(url, headers=headers, json=payload, timeout=10)
        if res.status_code != 200:
            print(f"⚠️ LINE API Error ({res.status_code}): {res.text}")
    except Exception as e:
        print(f"❌ LINE Reply Exception: {e}")

def create_line_flex_bubble(title, subtitle, fields, color="#1DB446"):
    """Creates a beautiful LINE Flex Message JSON bubble."""
    contents = []
    
    # Title/Header
    contents.append({
        "type": "text",
        "text": title,
        "weight": "bold",
        "size": "xl",
        "color": color
    })
    
    # Subtitle
    if subtitle:
        contents.append({
            "type": "text",
            "text": subtitle,
            "size": "sm",
            "color": "#aaaaaa",
            "wrap": True,
            "margin": "md"
        })
    
    contents.append({"type": "separator", "margin": "lg"})
    
    # Fields (key-value pairs)
    field_rows = []
    for key, val in fields.items():
        field_rows.append({
            "type": "box",
            "layout": "horizontal",
            "contents": [
                {
                    "type": "text",
                    "text": key,
                    "size": "sm",
                    "color": "#555555",
                    "flex": 1
                },
                {
                    "type": "text",
                    "text": str(val),
                    "size": "sm",
                    "color": "#111111",
                    "flex": 2,
                    "wrap": True
                }
            ],
            "margin": "md"
        })
    
    contents.append({
        "type": "box",
        "layout": "vertical",
        "contents": field_rows,
        "margin": "lg"
    })
    
    contents.append({"type": "separator", "margin": "lg"})
    
    # Footer
    contents.append({
        "type": "box",
        "layout": "horizontal",
        "contents": [
            {
                "type": "text",
                "text": "OrgChat Smart Helper",
                "size": "xs",
                "color": "#aaaaaa",
                "align": "end",
                "style": "italic"
            }
        ],
        "margin": "md"
    })

    return {
        "type": "bubble",
        "body": {
            "type": "box",
            "layout": "vertical",
            "contents": contents
        }
    }

def send_line_push_notification(target_username, title, text, fields=None):
    """Sends a one-to-one message to a specific user via LINE using Flex Messages where possible."""
    line_id = database.get_line_id_by_username(target_username)
    if not line_id:
        return
        
    token = os.environ.get("LINE_CHANNEL_ACCESS_TOKEN", "").strip()
    url = "https://api.line.me/v2/bot/message/push"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {token}"
    }
    
    if fields:
        # Use Flex Message
        flex_contents = create_line_flex_bubble(title, "การแจ้งเตือนส่วนตัว", fields, color="#E67E22")
        message_payload = {
            "type": "flex",
            "altText": f"แจ้งเตือน: {title}",
            "contents": flex_contents
        }
    else:
        # Fallback to text
        message_payload = {
            "type": "text", 
            "text": f"🔔 {title}\n{text}"
        }
        
    payload = {
        "to": line_id,
        "messages": [message_payload]
    }
    
    try:
        res = requests.post(url, headers=headers, json=payload, timeout=5)
        if res.status_code == 200:
            log_bot(f"📲 LINE Push sent to {target_username}")
    except Exception as e:
        print(f"❌ LINE Push Error: {e}")

def broadcast_line_announcement(title, text, fields=None):
    """Sends a message to all users who follow the bot using Flex Messages."""
    token = os.environ.get("LINE_CHANNEL_ACCESS_TOKEN", "").strip()
    if not token: return
    
    url = "https://api.line.me/v2/bot/message/broadcast"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {token}"
    }
    
    if fields:
        flex_contents = create_line_flex_bubble(title, "ประกาศองค์กร", fields, color="#007bff")
        message_payload = {
            "type": "flex",
            "altText": f"ประกาศ: {title}",
            "contents": flex_contents
        }
    else:
        # Fallback to text with nicer formatting
        message_payload = {
            "type": "text", 
            "text": f"📢 {title}\n━━━━━━━━━━━━━━\n{text[:512]}\n━━━━━━━━━━━━━━"
        }
        
    payload = {
        "messages": [message_payload]
    }
    try:
        res = requests.post(url, headers=headers, json=payload, timeout=10)
        return res.status_code == 200
    except Exception as e:
        print(f"❌ LINE Broadcast Error: {e}")
        return False

@app.route("/api/line/webhook", methods=["POST"])
def line_webhook():
    # อ่าน raw body ก่อนสิ่งอื่น (สำคัญ: ต้องอ่านก่อน request.json)
    body = request.get_data()
    signature = request.headers.get("X-Line-Signature", "").strip()

    if not signature:
        print("⚠️ [LINE] ไม่มี X-Line-Signature header มาด้วย", flush=True)
        return "No signature", 400

    if not verify_line_signature(body, signature):
        return "Invalid signature", 400
        
    data = request.json
    events = data.get("events", [])
    
    def process_line_event(event):
        if event.get("type") == "message" and event["message"].get("type") == "text":
            reply_token = event.get("replyToken")
            user_text = event["message"].get("text").strip()
            line_user_id = event["source"].get("userId")

            # --- Handler for Account Linking ---
            if user_text.lower().startswith("/link "):
                target_user = user_text[6:].strip()
                if database.link_line_user(target_user, line_user_id):
                    reply_to_line(reply_token, f"✅ ยินดีด้วยครับ! ชื่อบัญชี {target_user} ถูกเชื่อมต่อกับ LINE นี้แล้ว\nคุณจะได้รับการแจ้งเตือนสำคัญจาก OrgChat ทันทีครับ")
                else:
                    reply_to_line(reply_token, f"❌ ไม่พบผู้ใช้ '{target_user}' ในระบบ กรุณาตรวจสอบชื่อผู้ใช้ของคุณบนหน้าเว็บ OrgChat อีกครั้งครับ")
                return

            log_bot(f"📲 LINE Message: '{user_text[:30]}...' from {line_user_id[:8]}")
            context, sources = rag_engine.retrieve_context(user_text, where=None)
            activities = database.get_daily_activities()
            schedules = activities.get("schedules", [])
            posts = activities.get("posts", [])

            db_context = ""
            if schedules:
                sched_list = "\n".join([f"- {s['title']} วันที่ {s['date']} เวลา {s['time']} (หมวด: {s['category']})" for s in schedules])
                db_context += f"\n\n--- 📅 กำหนดการองค์กรล่าสุด ---\n{sched_list}"
            if posts:
                post_list = "\n".join([f"- {p['author']} โพสต์: {p['content'][:100]}..." for p in posts])
                db_context += f"\n\n--- 📢 โพสต์ข่าวสารล่าสุด ---\n{post_list}"
            if db_context:
                context += db_context

            weather_ctx = get_weather_context()
            system_prompt = (
                "คุณคือ 'AI-Assistant' ผู้ช่วยอัจฉริยะประจำองค์กรที่สื่อสารผ่าน LINE \n"
                "ข้อกำหนดในการตอบ:\n"
                "1. ตอบเป็นภาษาไทยที่สุภาพ เป็นมิตร และสรุปให้กระชับเหมาะสมกับการอ่านบนมือถือ\n"
                "2. ใช้ภาษาที่เป็นธรรมชาติ หลีกเลี่ยงสำนวนที่ดูเหมือนการแปลจากภาษาอื่น\n"
                "3. ลงท้ายด้วย 'ครับ' และใช้สรรพนาม 'ผม' หรือ 'เรา' ตามความเหมาะสม\n"
                "4. หากไม่พบข้อมูลใน Context ให้แจ้งอย่างสุภาพว่าไม่พบข้อมูลในฐานความรู้ปัจจุบัน\n"
                "5. หากผู้ใช้ถามเรื่องอากาศ อุณหภูมิ ฝนตก PM2.5 หรือ UV ให้ใช้ข้อมูลจากส่วน 'สภาพอากาศ' ด้านล่าง\n\n"
                f"{weather_ctx}\n"
                f"Context:\n{context}"
            )

            try:
                provider = ai_providers.get_provider()
                full_answer = ""
                for chunk in provider.chat_stream(user_text, [], system_prompt):
                    if chunk: full_answer += chunk

                if not full_answer:
                    full_answer = "ขออภัยครับ ผมไม่พบข้อมูลที่เกี่ยวข้องในระบบคลังความรู้ขององค์กร"

                if len(full_answer) > 4000:
                    full_answer = full_answer[:3900] + "\n...(มีเนื้อหาเพิ่มเติมในระบบเว็บ)..."

                reply_to_line(reply_token, full_answer)
                log_bot(f"✅ LINE Reply sent ({len(full_answer)} chars)")

            except Exception as e:
                print(f"❌ LINE AI Processing Error: {e}")
                reply_to_line(reply_token, "ขออภัยครับ เกิดข้อผิดพลาดในการประมวลผลข้อมูล")

    for event in events:
        threading.Thread(target=process_line_event, args=(event,), daemon=True).start()

    return "OK", 200


@app.route("/api/logout", methods=["POST"])
def api_logout():
    session.pop("user", None)
    return jsonify({"ok": True})

@app.route("/api/me")
def api_me():
    user = session.get("user")
    if user:
        profile = database.get_user_profile(user)
        user_settings = database.get_user_setting(user)
        return jsonify({
            "ok": True, 
            "user": user, 
            "profile": {
                **profile,
                "role": user_settings.get("role"),
                "can_edit_kb": user_settings.get("can_edit_kb", False),
                "is_active": user_settings.get("is_active", 1)
            }
        })
    return jsonify({"ok": False, "error": "Unauthorized"}), 401

@app.route("/api/profile/update", methods=["POST"])
@login_required
def api_update_profile():
    user = session.get("user")
    display_name = request.form.get("display_name")
    background_url = request.form.get("background_url")
    department = request.form.get("department")
    
    avatar_url = None
    if 'avatar' in request.files:
        file = request.files['avatar']
        if file.filename:
            # Create profiles dir if not exists
            prof_dir = Path("uploads/profiles")
            prof_dir.mkdir(parents=True, exist_ok=True)
            
            filename = secure_filename(f"{user}_{file.filename}")
            save_path = prof_dir / filename
            file.save(str(save_path))
            avatar_url = f"/uploads/profiles/{filename}"
            
    database.update_user_profile(user, display_name=display_name, avatar_url=avatar_url, background_url=background_url, department=department)
    database.log_event(f"Profile updated for {user}", user=user)
    return jsonify({"ok": True, "profile": database.get_user_profile(user)})

@app.route("/")
def index():
    google_client_id = os.environ.get("GOOGLE_CLIENT_ID", "")
    return render_template("index.html", google_client_id=google_client_id, vapid_public_key=VAPID_PUBLIC_KEY)


@app.route("/api/status")
@login_required
def status():
    provider = os.environ.get("AI_PROVIDER", "gemini").lower()
    has_key = False
    if provider == "gemini":
        has_key = bool(_get_gemini_api_key())
    elif provider == "groq":
        has_key = bool(os.environ.get("GROQ_API_KEY"))
    elif provider == "ollama":
        has_key = True # Local
        
    stats = rag_engine.kb_stats()
    quota_info = rag_engine.get_quota_status()
    
    user_settings = database.get_user_setting(session.get("user"))
    
    return jsonify({
        "api_key_set": has_key, 
        "quota_info": quota_info,
        "provider": provider,
        "is_admin": user_settings.get("role") == "admin",
        "can_view_kb": bool(user_settings.get("can_view_kb")),
        "can_edit_kb": bool(user_settings.get("can_edit_kb")),
        "can_delete_kb": bool(user_settings.get("can_delete_kb")),
        "user": session.get("user"),
        "app_settings": database.get_all_app_settings(),
        "server_time": get_current_time(),
        **stats
    })


@app.route("/api/set_key", methods=["POST"])
@admin_required
def set_key():
    data = request.get_json(force=True)
    key = data.get("api_key", "").strip()
    if not key:
        return jsonify({"ok": False, "error": "กรุณาใส่ API Key"}), 400
    os.environ["GEMINI_API_KEY"] = key
    # Persist to .env using absolute path from rag_engine
    env_path = rag_engine.BASE_DIR / ".env"
    lines = env_path.read_text(encoding="utf-8").splitlines() if env_path.exists() else []
    lines = [l for l in lines if not l.startswith("GEMINI_API_KEY")]
    lines.append(f"GEMINI_API_KEY={key}")
    env_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    
    # Refresh RAG engine with new key
    rag_engine._kb.update_api_key(key)
    
    return jsonify({"ok": True})


@app.route("/api/upload", methods=["POST"])
@login_required
def upload():
    if not can_edit_knowledge_base():
        return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์ในการอัปโหลดไฟล์ (ปิดการใช้งานโดยผู้ดูแลระบบ)"}), 403
    user = session.get("user", "Admin")
    import uuid
    files = request.files.getlist("files")
    dept = request.form.get("department", "General")
    category_id = request.form.get("category_id")
    
    # Convert category_id to int if present
    if category_id and category_id.isdigit():
        category_id = int(category_id)
    else:
        category_id = None

    if not files:
        return jsonify({"ok": False, "error": "ไม่พบไฟล์"}), 400

    results = []
    for f in files:
        orig_name = f.filename or "file"
        uid = str(uuid.uuid4())[:8]
        safe_name = secure_filename(orig_name)
        if not safe_name or safe_name.startswith('.'):
            safe_name = f"upload_{uid}"
            
        save_path = rag_engine.UPLOAD_DIR / f"{uid}_{safe_name}"
        f.save(str(save_path))

        # Use a thread for ingestion to keep the app responsive
        def run_ingest(path, name, d, cat):
            try:
                rag_engine.ingest_file(path, original_name=name, department=d, category_id=cat)
                # --- Notify LINE when KB is updated ---
                broadcast_line_announcement(
                    "คลังความรู้อัปเดต", 
                    f"-> เพิ่มไฟล์ใหม่เข้าระบบ!\nไฟล์: {name}\nส่วนงาน: {d}\nโดย: {user or 'Admin'}",
                    fields={
                        "ไฟล์": name,
                        "ส่วนงาน": d,
                        "ผู้บันทึก": user or "Admin"
                    }
                )
            except Exception as e:
                print(f"❌ Background Ingest Error: {e}")

        thread = threading.Thread(target=run_ingest, args=(save_path, orig_name, dept, category_id))
        thread.start()
        
        results.append({"file": orig_name, "status": "processing"})
        database.log_event(f"Started uploading file: {orig_name}", user=user)

    return jsonify({"ok": True, "results": results, "msg": "ไฟล์กำลังถูกประมวลผลในพื้นหลัง"})


@app.route("/api/files")
@login_required
def list_files():
    if not can_view_knowledge_base():
        return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์ในการเข้าถึงคลังข้อมูล"}), 403
    
    user = session.get("user", "Admin")
    all_files = rag_engine.list_files()
    
    # Get categories visible to this user
    visible_categories = database.get_categories(user)
    visible_cat_ids = {str(c["id"]) for c in visible_categories}
    
    # Filter files: visible if category is visible OR if file has no category
    filtered_files = []
    for f in all_files:
        cat_id = str(f.get("category_id") or "")
        if not cat_id or cat_id in visible_cat_ids:
            filtered_files.append(f)
            
    return jsonify({"files": filtered_files})

@app.route("/api/files/<file_id>", methods=["DELETE"])
@login_required
def delete_file_route(file_id):
    if not can_delete_knowledge_base():
        return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์ในการลบไฟล์ (ติดต่อผู้ดูแลระบบเพื่อขอสิทธิ์)"}), 403
    print(f"🗑️ Deletion request for file_id: {file_id}")
    ok = rag_engine.delete_file(file_id)
    if ok:
        user = session.get("user", "Admin")
        database.log_event(f"Deleted file ID: {file_id}", user=user)
        return jsonify({"ok": True})
    print(f"❌ Deletion failed for file_id: {file_id}")
    return jsonify({"ok": False, "error": "File not found or could not be deleted"}), 404

@app.route("/api/departments")
@login_required
def list_all_departments():
    conn = sqlite3.connect(database.DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT DISTINCT department FROM user_profiles WHERE department IS NOT NULL AND department != ''")
    rows = cursor.fetchall()
    conn.close()
    depts = [r[0] for r in rows]
    if "General" not in depts: depts.append("General")
    return jsonify({"ok": True, "departments": sorted(depts)})

# --- KB Category Routes ---

@app.route("/api/kb/categories", methods=["GET"])
@login_required
def list_categories():
    user = session.get("user", "Admin")
    categories = database.get_categories(user)
    return jsonify({"ok": True, "categories": categories})

@app.route("/api/kb/categories", methods=["POST"])
@login_required
def add_kb_category():
    if not can_edit_knowledge_base():
         return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์ในการจัดการหมวดหมู่"}), 403
    data = request.json
    name = data.get("name")
    desc = data.get("description")
    visibility = data.get("visibility", "public")
    user = session.get("user", "Admin")
    
    if not name:
        return jsonify({"ok": False, "error": "กรุณาระบุชื่อหมวดหมู่"}), 400
        
    try:
        database.add_category(name, desc, user, visibility)
        database.log_event(f"Added KB category: {name} (Visibility: {visibility})", user=user)
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route("/api/kb/categories/settings", methods=["POST"])
@login_required
def update_kb_category_settings():
    if not can_edit_knowledge_base():
         return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์ในการจัดการหมวดหมู่"}), 403
    data = request.json
    cat_id = data.get("id")
    visibility = data.get("visibility")
    authorized_users = data.get("authorized_users", []) # List of usernames
    user = session.get("user", "Admin")
    
    if not cat_id or not visibility:
        return jsonify({"ok": False, "error": "ข้อมูลไม่ครบถ้วน"}), 400
        
    # Check if user has permission to edit THIS category (must be owner or Admin)
    cat_info = next((c for c in database.get_categories("Admin") if c["id"] == cat_id), None)
    if not cat_info:
        return jsonify({"ok": False, "error": "ไม่พบหมวดหมู่"}), 404
        
    if user != "Admin" and cat_info["created_by"] != user:
        return jsonify({"ok": False, "error": "คุณไม่ใช่เจ้าของหมวดหมู่นี้"}), 403

    try:
        success = database.update_category_settings(cat_id, visibility, authorized_users)
        if success:
            database.log_event(f"Updated KB category settings ID: {cat_id}", user=user)
            return jsonify({"ok": True})
        return jsonify({"ok": False, "error": "ล้มเหลวในการบันทึกการตั้งค่า"}), 500
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route("/api/kb/categories/<int:cat_id>", methods=["DELETE"])
@login_required
def delete_kb_category(cat_id):
    if not can_edit_knowledge_base():
         return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์ในการจัดการหมวดหมู่"}), 403
    user = session.get("user", "Admin")
    try:
        database.delete_category(cat_id)
        database.log_event(f"Deleted KB category ID: {cat_id}", user=user)
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route("/api/kb/categories/access/<int:cat_id>", methods=["GET"])
@login_required
def get_kb_category_access(cat_id):
    # Only owner or Admin can see the access list
    user = session.get("user", "Admin")
    cat_info = next((c for c in database.get_categories("Admin") if c["id"] == cat_id), None)
    if not cat_info:
        return jsonify({"ok": False, "error": "ไม่พบหมวดหมู่"}), 404
        
    if user != "Admin" and cat_info["created_by"] != user:
        return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์เข้าถึงข้อมูลนี้"}), 403

    conn = sqlite3.connect(database.DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT username FROM user_category_access WHERE category_id = ?", (cat_id,))
    users = [r[0] for r in cursor.fetchall()]
    conn.close()
    
    return jsonify({"ok": True, "users": users})

@app.route("/api/kb/search")
@login_required
def kb_search():
    if not can_view_knowledge_base():
        return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์ในการค้นหาคลังข้อมูล"}), 403
    """Search knowledge base files by name or content snippet."""
    query = request.args.get("q", "").strip().lower()
    if not query:
        return jsonify({"ok": True, "results": []})
        
    user = session.get("user", "Admin")
    visible_categories = database.get_categories(user)
    visible_cat_ids = [str(c["id"]) for c in visible_categories]
    
    # ChromaDB 'where' filter logic
    if visible_cat_ids:
        # User sees: restricted/private they own + public + unassigned
        where_filter = {
            "$or": [
                {"category_id": {"$in": visible_cat_ids}},
                {"category_id": ""} # Unassigned
            ]
        }
    else:
        # Fallback: only unassigned or public if user has no visible cats?
        # Actually database.get_categories(user) always returns public ones.
        where_filter = {"category_id": ""}

    try:
        files = rag_engine.list_files()
        # Filter by file name (fast, local), respecting visibility
        name_matches = []
        for f in files:
            cat_id = str(f.get("category_id") or "")
            if not cat_id or cat_id in visible_cat_ids:
                if query in (f.get("name") or "").lower():
                    name_matches.append(f)
                    
        # Also do a semantic search for content matches
        kb_results = rag_engine.search_kb(query, n_results=10, where=where_filter)
        
        # Combine: file cards for name matches + content snippets for semantic
        results = []
        seen_ids = set()
        for f in name_matches:
            results.append({
                "file_id": f["file_id"],
                "file_name": f["name"],
                "file_type": f.get("type", ""),
                "text": f"ชื่อไฟล์ตรงกับคำค้นหา",
                "score": 1.0,
                "location": f.get("department", "General")
            })
            seen_ids.add(f["file_id"])
        
        for r in kb_results:
            fid = r.get("file_id", "")
            if fid not in seen_ids:
                results.append({
                    "file_id": fid,
                    "file_name": r.get("name", r.get("source", "")),
                    "file_type": r.get("file_type", ""),
                    "text": r.get("text", "")[:200],
                    "score": r.get("score", 0),
                    "location": r.get("department", "General")
                })
                seen_ids.add(fid)
        return jsonify({"ok": True, "results": results[:20]})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route("/api/kb/files/assign", methods=["POST"])
@login_required
def assign_file_category():
    if not can_edit_knowledge_base():
         return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์ในการจัดการไฟล์"}), 403
    data = request.json
    file_id = data.get("file_id")
    category_id = data.get("category_id") # Can be None to unassign
    user = session.get("user", "Admin")

    if not file_id:
        return jsonify({"ok": False, "error": "Missing file_id"}), 400

    try:
        # Update in database if needed, but primarily in rag_engine meta
        rag_engine.update_file_category(file_id, category_id)
        database.log_event(f"Assigned file {file_id} to category {category_id}", user=user)
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
@app.route("/api/personas", methods=["GET"])
@login_required
def get_personas():
    personas = database.get_all_personas_v2()
    return jsonify({"ok": True, "personas": personas})

@app.route("/api/kb/files/view/<file_id>")
@login_required
def view_kb_file(file_id):
    if not can_view_knowledge_base():
        return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์ในการดูไฟล์"}), 403
    files = rag_engine.list_files()
    file_info = next((f for f in files if f["file_id"] == file_id), None)
    if not file_info:
        return "File not found", 404
    
    path = Path(file_info["path"])
    if not path.exists():
        print(f"❌ File not found on disk: {path}")
        return f"File not found on disk: {path}", 404
        
    return send_file(path, mimetype='application/pdf' if file_info['type'] == 'pdf' else 'text/plain')

@app.route("/api/export/pdf", methods=["POST"])
@login_required
def export_pdf():
    try:
        data = request.get_json(force=True)
        title = data.get("title", "OrgChat Report")
        content = data.get("content", "")
        
        if not content:
            return jsonify({"ok": False, "error": "ไม่มีเนื้อหาให้ส่งออก"}), 400

        pdf = FPDF()
        pdf.add_page()
        
        # We need a Thai font. For now, we'll try to use a standard one and add Thai support once font path is confirmed.
        # pdf.add_font('THSarabunNew', '', 'font_path', unicode=True)
        # pdf.set_font('THSarabunNew', '', 16)
        
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt=title, ln=1, align='C')
        pdf.ln(10)
        
        # Multi_cell handles word wrap
        pdf.multi_cell(0, 10, txt=content)
        
        filename = f"report_{uuid.uuid4().hex[:8]}.pdf"
        output_path = rag_engine.UPLOAD_DIR / filename
        pdf.output(str(output_path))
        
        return jsonify({
            "ok": True, 
            "url": f"/api/files/download/{filename}",
            "filename": filename
        })
    except Exception as e:
        print(f"❌ PDF Export Error: {e}")
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/files/download/<filename>")
def download_file(filename):
    return send_from_directory(rag_engine.UPLOAD_DIR, filename, as_attachment=True)

# --- Calendar Schedules ---
@app.route("/api/schedules", methods=["GET", "POST"])
@login_required
def manage_schedules():
    user = session.get("user")
    if request.method == "POST":
        data = request.get_json(force=True)
        title = data.get("title")
        start_date = data.get("date") # Expected YYYY-MM-DD
        desc = data.get("desc", "")
        category = data.get("category", "General")
        start_time = data.get("time", "09:00")
        is_public = 1 if data.get("is_public") else 0
        target_depts = data.get("target_departments")
        target_users = data.get("target_users")
        status = data.get("status", "todo")
        
        if not title or not start_date:
            return jsonify({"ok": False, "error": "Missing title or date"}), 400
            
        database.add_schedule(user, title, start_date, desc, category, start_time, is_public, status, target_departments=target_depts, target_users=target_users)
        visibility_text = "Public" if is_public else "Private"
        database.log_event(f"Added {visibility_text} schedule: {title} on {start_date}", user=user)
        
        # --- New Public Event Broadcast ---
        profile = database.get_user_profile(user)
        display_name = profile.get("display_name", user)
        if is_public:
            notification_db.notify_all_except(
                user, 
                'calendar', 
                'กิจกรรมใหม่: ' + title, 
                f'โดย {display_name} | วันที่ {start_date} เวลา {start_time}',
                link='#calendar'
            )
            # --- LINE Broadcast for Public Events ---
            threading.Thread(target=broadcast_line_announcement, args=(
                "กิจกรรมใหม่", 
                f"-> กำหนดการใหม่!\nหัวข้อ: {title}\nวันที่: {start_date}\nเวลา: {start_time} น.\nโดย: {display_name}",
            ), kwargs={
                "fields": {
                    "หัวข้อ": title,
                    "วันที่": start_date,
                    "เวลา": f"{start_time} น.",
                    "โดย": display_name
                }
            }).start()
        else:
            # Notify specific users if specified
            if target_users:
                t_users = [u.strip() for u in target_users.split(',') if u.strip() and u.strip().lower() != user.lower()]
                if t_users:
                    notification_db.notify_users(
                        t_users, 'calendar', 'กำหนดการใหม่ที่แชร์กับคุณ: ' + title,
                        f'โดย {display_name} | วันที่ {start_date} เวลา {start_time}',
                        link='#calendar'
                    )
                    batch_send_push_notification(t_users, 'กำหนดการที่แชร์กับคุณ', f'{display_name}: {title}', url='#calendar')
                    
                    # --- LINE Push for Targeted Users ---
                    for target in t_users:
                        threading.Thread(target=send_line_push_notification, args=(
                            target, 
                            "กิจกรรมใหม่", 
                            f"-> กำหนดการใหม่สำหรับคุณ!\nหัวข้อ: {title}\nวันที่: {start_date}\nเวลา: {start_time} น.",
                        ), kwargs={
                            "fields": {
                                "หัวข้อ": title,
                                "วันที่": start_date,
                                "เวลา": f"{start_time} น.",
                                "โดย": display_name
                            }
                        }).start()
            
            # Notify specific departments
            if target_depts:
                t_depts = [d.strip() for d in target_depts.split(',') if d.strip()]
                # Get all users in these departments
                conn = sqlite3.connect(database.DB_PATH)
                cursor = conn.cursor()
                cursor.execute(f"SELECT username FROM user_profiles WHERE department IN ({','.join(['?']*len(t_depts))})", t_depts)
                dept_users = [r[0] for r in cursor.fetchall() if r[0].lower() != user.lower()]
                conn.close()
                if dept_users:
                    notification_db.notify_users(
                        dept_users, 'calendar', 'กิจกรรมใหม่ในแผนก: ' + title,
                        f'โดย {display_name} | วันที่ {start_date} เวลา {start_time}',
                        link='#calendar'
                    )
                    batch_send_push_notification(dept_users, 'กิจกรรมใหม่ในแผนก', f'{display_name}: {title}', url='#calendar')
                        
        return jsonify({"ok": True})
    
    return jsonify({"schedules": database.get_schedules(user)})

@app.route("/api/schedules/<int:sid>", methods=["PUT", "DELETE"])
@login_required
def manage_schedule_item(sid):
    user = session.get("user")
    if request.method == "DELETE":
        ok = database.delete_schedule(sid, username=user, is_admin=is_admin())
        if ok:
            database.log_event(f"Schedule deleted: ID {sid}", user=user)
            return jsonify({"ok": True})
        return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์ลบกำหนดการนี้"}), 403
    
    if request.method == "PUT":
        data = request.get_json(force=True)
        title = data.get("title")
        date = data.get("date")
        desc = data.get("desc", "")
        cat = data.get("category", "General")
        time_val = data.get("time", "09:00")
        is_public = 1 if data.get("is_public") else 0
        status = data.get("status", "todo")
        target_depts = data.get("target_departments")
        target_users = data.get("target_users")
        
        if not title or not date:
            return jsonify({"ok": False, "error": "Missing title or date"}), 400
            
        database.update_schedule(sid, title, date, desc, cat, time_val, is_public, status, target_departments=target_depts, target_users=target_users)
        user = session.get("user", "Admin")
        database.log_event(f"Updated schedule ID: {sid}", user=user)
        return jsonify({"ok": True})

@app.route("/api/schedules/<int:sid>/toggle", methods=["POST"])
@login_required
def toggle_schedule_route(sid):
    new_status = database.toggle_schedule_status(sid)
    if new_status is None:
        return jsonify({"ok": False, "error": "Schedule not found"}), 404
    user = session.get("user", "Admin")
    database.log_event(f"Toggled schedule status: {sid} to {new_status}", user=user)
    return jsonify({"ok": True, "new_status": new_status})

@app.route("/api/schedules/clear-past", methods=["DELETE"])
@login_required
def delete_past_schedules_route():
    user = session.get("user")
    database.delete_past_schedules(user)
    database.log_event(f"Cleared all past schedules", user=user)
    return jsonify({"ok": True})

@app.route("/api/schedules/archive-past", methods=["POST"])
@login_required
def archive_past_schedules_route():
    user = session.get("user")
    database.archive_past_schedules(user)
    database.log_event(f"Archived all past schedules", user=user)
    return jsonify({"ok": True})



@app.route("/api/ping")
def ping():
    return jsonify({"ok": True, "version": "v1.2-diagnostics"})

# ⚡ AI Kanban Auto-Generator
@app.route("/api/kanban/auto-generate", methods=["POST"])
@login_required
def auto_generate_kanban():
    user = session.get("user", "Admin")
    data = request.get_json(force=True)
    goal = data.get("goal")
    if not goal:
        return jsonify({"ok": False, "error": "กรุณากรอกเป้าหมาย"}), 400

    prompt = (
        f"จากเป้าหมายดังนี้: '{goal}' "
        "ช่วยแตกเป็นรายการงานย่อยๆ (Tasks) สำหรับระบบ Kanban "
        "ตอบในรูปแบบ JSON เป็นรายการดั้งนี้ "
        '{"tasks": [{"title": "ชื่อสั้นๆ", "desc": "รายละเอียดงาน", "category": "General|Task|Meeting", "date": "YYYY-MM-DD", "time": "HH:MM"}]} '
        f"กำหนดวันที่เริ่มจากวันนี้ ({get_current_time().split(' ')[0]}) เป็นต้นไป "
        "ข้อสำคัญ: ตอบเฉพาะ JSON เท่านั้น ห้ามมีข้อความอื่นเด็ดขาด"
    )

    try:
        provider = ai_providers.get_provider()
        response_text = provider.chat(prompt, [], "")
        
        # Clean potential markdown backticks
        clean_json = response_text.replace("```json", "").replace("```", "").strip()
        result = json.loads(clean_json)
        
        tasks_added = 0
        for t in result.get("tasks", []):
            database.add_schedule(
                user, 
                t.get("title"), 
                t.get("date"), 
                description=t.get("desc", ""), 
                category=t.get("category", "Task"),
                time=t.get("time", "09:00"),
                status="todo"
            )
            tasks_added += 1
            
        database.log_event(f"AI Auto-generated {tasks_added} tasks for goal: {goal[:50]}", user=user)
        return jsonify({"ok": True, "count": tasks_added})
    except Exception as e:
        print(f"❌ AI Kanban Error: {e}")
        return jsonify({"ok": False, "error": str(e)}), 500

# 📑 AI Document Comparison
@app.route("/api/kb/compare", methods=["POST"])
@login_required
def compare_knowledge_base_files():
    if not can_view_knowledge_base():
        return jsonify({"ok": False, "error": "Unauthorized"}), 403
        
    data = request.json
    f1_id = data.get("file1_id")
    f2_id = data.get("file2_id")
    prompt_custom = data.get("prompt", "ช่วยสรุปสิ่งที่เหมือนกัน, สิ่งที่แตกต่างกัน และจุดสำคัญที่คุณค้นพบ")
    
    if not f1_id or not f2_id:
        return jsonify({"ok": False, "error": "Missing file IDs"}), 400
        
    try:
        # Retrieve ALL content for both files
        ctx1 = rag_engine.get_file_content(f1_id)
        ctx2 = rag_engine.get_file_content(f2_id)
        
        if not ctx1 or not ctx2:
            return jsonify({"ok": False, "error": "ไม่สามารถดึงข้อมูลเนื้อหาไฟล์มาเปรียบเทียบได้"}), 404

        prompt = (
            f"ช่วยเปรียบเทียบข้อมูลจากสองแหล่งนี้อย่างละเอียด:\n\n"
            f"--- ข้อมูลชุดที่ 1 ---\n{ctx1[:8000]}\n\n"
            f"--- ข้อมูลชุดที่ 2 ---\n{ctx2[:8000]}\n\n"
            f"คำสั่งเพิ่มเติม: {prompt_custom}\n\n"
            "ตอบเป็นภาษาไทยในรูปแบบ Markdown ที่สวยงาม มีตารางเปรียบเทียบ (ถ้าเป็นไปได้) และสรุปประเด็นสำคัญสำหรับผู้บริหาร"
        )
        
        provider = ai_providers.get_provider()
        reply = provider.chat_stream(prompt, [], "คุณคือนักวิเคราะห์ข้อมูลผู้ช่วยส่วนกลางขององค์กร")
        
        # We'll use a simple non-streaming response for compare for now to keep it simple, or we could stream it.
        # But for comparison, a solid final result is often preferred.
        full_reply = ""
        for chunk in reply:
             if chunk: full_reply += chunk
             
        return jsonify({"ok": True, "comparison": full_reply})
    except Exception as e:
        print(f"❌ Document Comparison Error: {e}")
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/csv/<file_id>", methods=["GET", "POST"])
@login_required
def manage_csv_combined(file_id):
    """Combined CSV handler to ensure registration."""
    print(f"🚀 CSV API HIT: {request.method} {file_id}")
    files = rag_engine.list_files()
    target = next((f for f in files if f["file_id"] == file_id), None)
    if not target:
        return jsonify({"ok": False, "error": "File not found"}), 404
        
    path = Path(target["path"])
    
    if request.method == "GET":
        if not can_view_knowledge_base():
            return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์ในการเข้าถึงไฟล์"}), 403
        if not path.exists() or path.suffix.lower() != ".csv":
            return jsonify({"ok": False, "error": "Invalid file type"}), 400
        try:
            import csv
            data = []
            with open(path, newline='', encoding='utf-8-sig') as f:
                reader = csv.DictReader(f)
                headers = reader.fieldnames
                for row in reader:
                    data.append(row)
            return jsonify({"ok": True, "headers": headers, "data": data, "name": target["name"]})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    if request.method == "POST":
        if not can_edit_knowledge_base():
            return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์ในการแก้ไขไฟล์ (ปิดการใช้งานโดยผู้ดูแลระบบ)"}), 403
        data_json = request.get_json()
        rows = data_json.get("data", [])
        headers = data_json.get("headers", [])
        if not rows or not headers:
            return jsonify({"ok": False, "error": "Empty data"}), 400
        try:
            import csv
            with open(path, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.DictWriter(f, fieldnames=headers)
                writer.writeheader()
                writer.writerows(rows)
            rag_engine.delete_file(file_id, delete_from_disk=False)
            res = rag_engine.ingest_file(path, original_name=target["name"], department=target.get("department", "General"))
            if res.get("status") == "error":
                return jsonify({"ok": False, "error": f"บันทึกไฟล์สำเร็จ แต่ AI อัปเดตข้อมูลล้มเหลว: {res.get('error')}"}), 500
            return jsonify({"ok": True})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/txt/<file_id>", methods=["GET", "POST"])
@login_required
def manage_txt(file_id):
    print(f"🚀 TXT API HIT: {request.method} {file_id}")
    files = rag_engine.list_files()
    target = next((f for f in files if f["file_id"] == file_id), None)
    if not target:
        return jsonify({"ok": False, "error": "File not found"}), 404
    path = Path(target["path"])
    if request.method == "GET":
        if not can_edit_knowledge_base():
            return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์ในการเข้าถึงตัวแก้ไขไฟล์"}), 403
        try:
            with open(path, "r", encoding="utf-8-sig") as f:
                content = f.read()
            return jsonify({"ok": True, "content": content})
        except Exception as e: return jsonify({"ok": False, "error": str(e)}), 500
    if request.method == "POST":
        if not can_edit_knowledge_base():
            return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์ในการแก้ไขไฟล์ (ปิดการใช้งานโดยผู้ดูแลระบบ)"}), 403
        try:
            content = request.json.get("content", "")
            with open(path, "w", encoding="utf-8") as f: f.write(content)
            rag_engine.delete_file(file_id, delete_from_disk=False)
            res = rag_engine.ingest_file(path, original_name=target["name"], department=target.get("department", "General"))
            if res.get("status") == "error":
                return jsonify({"ok": False, "error": f"บันทึกไฟล์สำเร็จ แต่ AI อัปเดตข้อมูลล้มเหลว: {res.get('error')}"}), 500
            return jsonify({"ok": True})
        except Exception as e: return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/search")
@login_required
def search():
    query = request.args.get("q", "").strip()
    if not query:
        return jsonify({"results": []})
    
    user = session.get("user")
    rag_filter = get_rag_filter(user)
    results = rag_engine.search_kb(query, n_results=10, where=rag_filter)
    return jsonify({"results": results})


@app.route("/api/sync", methods=["POST"])
@admin_required
def sync_kb():
    results = rag_engine.sync_uploads()
    return jsonify({"ok": True, "results": results})

@app.route("/api/prune", methods=["POST"])
@admin_required
def prune_kb_route():
    count = rag_engine.prune_kb()
    return jsonify({"ok": True, "pruned": count})


@app.route("/api/wipe", methods=["POST"])
@admin_required
def wipe():
    ok = rag_engine.wipe_knowledge_base()
    return jsonify({"ok": ok})


@app.route("/api/history")
@login_required
def get_history():
    user = session.get("user")
    return jsonify({"history": database.get_history(user)})



@app.route("/api/history/clear", methods=["POST"])
@login_required
def clear_history():
    user = session.get("user")
    database.clear_history(user)
    return jsonify({"ok": True})


@app.route("/api/feedback", methods=["POST"])
def feedback():
    user = session.get("user", "Admin")
    data = request.get_json(force=True)
    val = data.get("value", 0) # 1 or -1
    database.save_feedback(user, val)
    return jsonify({"ok": True})


@app.route("/api/stats")
@login_required
def stats_route():
    user = session.get("user")
    return jsonify(database.get_stats(user))

@app.route("/api/dashboard/data")
@login_required
def dashboard_data():
    user = session.get("user")
    # Auto cleanup/archive very old data (90+ days)
    database.auto_archive_old_schedules(90)

    conn = sqlite3.connect(database.DB_PATH)
    cursor = conn.cursor()
    
    # 1. Total Files & Chunks
    kb_info = rag_engine.kb_stats()
    
    # 2. Total Chat Messages for this user
    cursor.execute("SELECT COUNT(*) FROM messages WHERE username = ?", (user,))
    chat_count = cursor.fetchone()[0]
    
    # 3. Total Users
    cursor.execute("SELECT COUNT(DISTINCT username) FROM user_profiles") # Simplified for now
    user_count = cursor.fetchone()[0]
    
    # 4. Recent Activity (Latest 5 events for this user)
    cursor.execute("SELECT event_text, time FROM events WHERE user = ? ORDER BY time DESC LIMIT 5", (user,))
    logs = cursor.fetchall()
    activity_count = len(logs)
    
    # 5. Recent History (Latest 5 chats)
    cursor.execute("SELECT text, timestamp FROM messages WHERE username = ? AND role = 'user' ORDER BY timestamp DESC LIMIT 5", (user,))
    recent_chats = [{"text": r[0], "time": r[1]} for r in cursor.fetchall()]
    
    # 7. Recent Files (Filtered by visibility)
    all_files = rag_engine.list_files()
    visible_categories = database.get_categories(user)
    visible_cat_ids = {str(c["id"]) for c in visible_categories}
    
    filtered_recent = []
    for f in all_files:
        cat_id = str(f.get("category_id") or "")
        if not cat_id or cat_id in visible_cat_ids:
            filtered_recent.append(f)
            
    recent_files = sorted(filtered_recent, key=lambda x: x.get('uploaded_at', ''), reverse=True)[:5]
    
    # 8. Unread Notifications Count
    import notification_db
    unread_notifs = notification_db.get_notifications(user, unread_only=True)
    unread_count = len(unread_notifs)

    # 9. Tasks and Schedules
    from datetime import date
    today = date.today().isoformat()
    user_schedules = database.get_schedules(user)
    pending_tasks = [s for s in user_schedules if s.get("status", "todo") in ("todo", "doing") and s.get("date", "") >= today]
    completed_tasks_count = len([s for s in user_schedules if s.get("status") == "done"])


    conn.close()
    
    return jsonify({
        "ok": True,
        "stats": {
            "files": len(filtered_recent), # Show only visible file count to user
            "chats": chat_count,
            "users": user_count,
            "activity": activity_count,
            "unread_notifications": unread_count,
            "completed_tasks": completed_tasks_count
        },
        "recent_chats": recent_chats,
        "recent_files": recent_files,
        "upcoming": pending_tasks[:3], 
        "pending_tasks": pending_tasks,
        "logs": [{"event": l[0], "time": l[1]} for l in logs]
    })

@app.route("/api/dashboard/briefing")
@login_required
def dashboard_briefing():
    user = session.get("user")
    activities = database.get_daily_activities()
    
    if not activities["posts"] and not activities["schedules"]:
        return jsonify({"ok": True, "briefing": "วันนี้ยังไม่มีกิจกรรมใหม่ที่บันทึกไว้ครับ คุณสามารถเริ่มแชทหรืออัปโหลดเอกสารใหม่เพื่อเริ่มต้นวันได้ทันที!"})
    
    # Prepare context for AI
    context = "กิจกรรมล่าสุดใน 24 ชั่วโมงที่ผ่านมา:\n"
    for p in activities["posts"][:5]:
        context += f"- [FEED] {p['author']} โพสต์ในหมวด {p['category']}: {p['content'][:100]}...\n"
    for s in activities["schedules"][:5]:
        context += f"- [SCHEDULE] มีนัดหมายถึง: {s['title']} วันที่ {s['date']} เวลา {s['time']}\n"
        
    prompt = f"""คุณเป็น AI ผู้ช่วยอัจฉริยะ (OrgChat) สรุป "Daily Briefing" สั้นๆ ให้กับผู้ใช้ชื่อ {user} 
จากข้อมูลกิจกรรมล่าสุดต่อไปนี้ โดยใช้ภาษาที่เป็นกันเอง ทันสมัย และกระตือรือร้น (ไม่เกิน 3-4 ประโยค):

{context}

สรุปให้น่าสนใจและเน้นสิ่งที่ควรทราบที่สุดประจำวันนี้:"""

    try:
        # ใช้ ai_providers ใหม่แทน genai.GenerativeModel เก่า
        provider = ai_providers.get_provider()
        briefing = ""
        for chunk in provider.chat_stream(prompt, [], "คุณคือ AI ผู้ช่วยสรุปข่าวสารองค์กร"):
            if chunk: briefing += chunk
        briefing = briefing.strip()
        return jsonify({"ok": True, "briefing": briefing})
    except Exception as e:
        print(f"Briefing Error: {e}")
        return jsonify({"ok": False, "error": "ไม่สามารถสร้างสรุปได้ในขณะนี้"})

@app.route("/api/summary/generate", methods=["POST"])
@login_required
def generate_global_summary():
    data = request.json or {}
    focus = data.get("focus", "").strip()
    category_id = data.get("category_id", "all")
    
    user = session.get("user", "Admin")
    visible_categories = database.get_categories(user)
    visible_cat_ids = [str(c["id"]) for c in visible_categories]
    
    # Global visibility filter (user sees visible categories + unassigned)
    global_where = {
        "$or": [
            {"category_id": {"$in": visible_cat_ids}},
            {"category_id": ""}
        ]
    }

    where_filter = None
    search_query = focus
    
    if category_id != "all":
        if category_id == "unassigned":
            where_filter = {"category_id": ""}
            if not search_query: search_query = "สรุปข้อมูลที่ยังไม่ได้ระบุหมวดหมู่"
        else:
            if str(category_id) not in visible_cat_ids:
                return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์เข้าถึงหมวดหมู่นี้"}), 403
            where_filter = {"category_id": str(category_id)}
            # Try to get category name for better context if no focus given
            if not search_query:
                try:
                    conn = sqlite3.connect(database.DB_PATH)
                    cursor = conn.cursor()
                    cursor.execute("SELECT name FROM kb_categories WHERE id = ?", (int(category_id),))
                    row = cursor.fetchone()
                    conn.close()
                    if row: search_query = f"สรุปข้อมูลเกี่ยวกับ {row[0]}"
                except: pass
    # else: category_id == "all" → where_filter stays None (search ALL data)
        
    if not search_query:
        search_query = "ภาพรวมองค์กร กฎระเบียบ ประกาศ ข่าวสารล่าสุด"

    print(f"📄 Summary Request - Category: {category_id}, Focus: '{focus}', Filter: {where_filter}")
    
    try:
        is_fallback = False
        # 1. Semantic Search
        results = rag_engine.search_kb(search_query, n_results=10, where=where_filter)
        
        # Fallback 1: Literal grab from category
        if not results:
            results = rag_engine.get_all_chunks(where=where_filter, limit=10)
            
        # Fallback 2: Try without any filter at all
        if not results and where_filter is not None:
            print(f"⚠️ Filtered search empty, falling back to unfiltered search")
            results = rag_engine.search_kb(search_query, n_results=10)
            if not results:
                results = rag_engine.get_all_chunks(limit=10)
            is_fallback = True

        if not results:
            msg = "ไม่มีข้อมูลในหมวดหมู่ที่ระบุหรือหมวดหมู่ที่คุณสามารถเข้าถึงได้"
            return jsonify({"ok": False, "error": msg}), 404
            
        context_text = "\n".join([f"- {r['text']}" for r in results])
        # Cap context for Groq TPM limits
        if len(context_text) > 8000:
            context_text = context_text[:8000] + "..."

        focus_instruction = f"\nเน้นข้อมูลเกี่ยวกับหัวข้อ: '{focus}'" if focus else ""
        fallback_notice = "\nหมายเหตุ: เนื่องจากหมวดหมู่ที่ระบุไม่มีข้อมูลเพียงพอ AI จึงใช้ข้อมูลจากคลังความรู้ที่คุณสามารถเข้าถึงได้ทั้งหมดในการสรุปแทน" if is_fallback else ""
        
        prompt = f"""คุณคือผู้ช่วยอัจฉริยะประจำบริษัท หน้าที่ของคุณคือสร้าง 'บทสรุปความรู้อันเป็นประโยชน์' จากข้อมูลที่ได้รับด้านล่างนี้
กรุณาสรุปด้วยภาษาไทยที่สุภาพ เป็นมืออาชีพ เน้นความกระชับและอ่านง่าย (ใช้ Markdown)
{focus_instruction}
{fallback_notice}

สำคัญ: หลีกเลี่ยงการแปลตรงตัวจากภาษาอังกฤษ ใช้สำนวนไทยที่คนทำงานออฟฟิศใช้กันจริงๆ และลงท้ายด้วย 'ครับ' อย่างเหมาะสม

โครงสร้างสรุป:
1. ภาพรวมหลัก (Main Overview)
2. ประเด็นสำคัญ / กฎระเบียบ (Key Points / Rules)
3. ข้อมูลอ้างอิงที่มีในระบบ (Sources Found)

ข้อมูลดิบจากคลังความรู้:
{context_text}"""

        provider = ai_providers.get_provider()
        summary_md = ""
        instruction = "สรุปข้อมูลบทสรุปความรู้องค์กรให้ฉันหน่อย"
        for chunk in provider.chat_stream(instruction, [], prompt):
            if chunk: summary_md += chunk
        
        if not summary_md:
            return jsonify({"ok": False, "error": "AI ไม่สามารถสร้างบทสรุปได้"}), 500
            
        keywords = list(set([r.get('department', 'General') for r in results]))
        refs = list(set([r.get('name', r.get('source', 'Unknown')) for r in results]))
        
        return jsonify({
            "ok": True, 
            "summary": summary_md,
            "keywords": keywords[:10],
            "refs": refs[:8],
            "is_fallback": is_fallback
        })
    except Exception as e:
        print(f"❌ Global Summary Error: {e}")
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/summarize_data", methods=["POST"])
@login_required
def summarize_data():
    """Generate an AI-powered summary/insight for CSV chart data shown on the viz page."""
    data = request.json or {}
    x_col = data.get("x_col", "X")
    y_col = data.get("y_col", "Y")
    data_context = data.get("data_context", "")

    if not data_context:
        return jsonify({"ok": False, "error": "ไม่มีข้อมูลให้วิเคราะห์"}), 400

    prompt = f"""คุณคือผู้เชี่ยวชาญด้านการวิเคราะห์ข้อมูลองค์กร หน้าที่ของคุณคือถอดรหัสตัวเลขให้เป็น 'ข้อมูลเชิงกลยุทธ์' ที่เข้าใจง่าย

ข้อมูลที่ต้องวิเคราะห์:
- แกน X (หมวดหมู่): "{x_col}"
- แกน Y (ค่าตัวเลข): "{y_col}"
- ข้อมูล: {data_context[:1200]}

สรุปแนวโน้มและจุดที่น่าสนใจเป็นภาษาไทยที่กระชับ ตรงไปตรงมา และมีสาระสำคัญที่นำไปใช้งานต่อได้
หลีกเลี่ยงศัพท์เทคนิคที่ซับซ้อนเกินไป และให้คำแนะนำที่ดูเป็นที่ปรึกษาที่เป็นมิตร

ตอบกระชับไม่เกิน 6 ประโยค ใช้ภาษาไทยที่เป็นธรรมชาติที่สุด"""

    try:
        # ใช้ ai_providers ใหม่แทน genai.GenerativeModel เก่า
        provider = ai_providers.get_provider()
        summary = ""
        for chunk in provider.chat_stream(prompt, [], "คุณคือนักวิเคราะห์ข้อมูลองค์กร"):
            if chunk: summary += chunk
        summary = summary.strip()
        return jsonify({"ok": True, "summary": summary})
    except Exception as e:
        print(f"❌ Summarize Data Error: {e}")
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route("/api/admin/dashboard/stats")
@admin_required
def admin_dashboard_stats():
    """Provides system-wide statistics for the admin dashboard."""
    conn = sqlite3.connect(database.DB_PATH)
    cursor = conn.cursor()
    
    # Total Queries (All users)
    cursor.execute("SELECT COUNT(*) FROM messages WHERE role = 'user'")
    total_queries = cursor.fetchone()[0]
    
    # Active Users (last 30 days)
    cursor.execute("SELECT COUNT(DISTINCT username) FROM user_profiles") # Simplified for now
    total_users = cursor.fetchone()[0]
    
    # Storage Stats
    uploads_size = sum(f.stat().st_size for f in rag_engine.UPLOAD_DIR.glob('*') if f.is_file())
    
    # Feedback counts
    cursor.execute("SELECT feedback, COUNT(*) FROM messages WHERE role = 'bot' GROUP BY feedback")
    feedback_stats = dict(cursor.fetchall())
    
    conn.close()
    
    return jsonify({
        "ok": True,
        "total_queries": total_queries,
        "total_users": total_users,
        "uploads_size_mb": round(uploads_size / (1024 * 1024), 2),
        "feedback": feedback_stats,
        "kb_size": rag_engine.kb_stats().get("knowledge_base_size", 0)
    })

# --- Admin Settings Management ---
@app.route("/api/admin/settings", methods=["GET", "POST"])
@admin_required
def admin_settings_route():
    if request.method == "GET":
        return jsonify({"ok": True, "settings": database.get_all_app_settings()})
    
    if request.method == "POST":
        data = request.json
        for key, value in data.items():
            database.set_app_setting(key, str(value))
        return jsonify({"ok": True})


@app.route("/api/admin/users", methods=["GET"])
@admin_required
def admin_list_users():
    users = database.admin_get_all_users()
    return jsonify({"ok": True, "users": users})

@app.route("/api/users/list", methods=["GET"])
@login_required
def list_users_for_sharing():
    """Returns a simplified list of all users for selection in sharing UI."""
    users = database.admin_get_all_users() # This returns profiles + settings
    # Filter only necessary fields for non-admin use
    safe_users = [{
        "username": u["username"],
        "display_name": u["display_name"] or u["username"].capitalize(),
        "avatar_url": u.get("avatar_url"),
        "department": u.get("department", "General")
    } for u in users if u.get("is_active", 1)]
    return jsonify({"ok": True, "users": safe_users})

@app.route("/api/admin/users", methods=["POST"])
@admin_required
def admin_create_user_route():
    data = request.json
    username = data.get("username", "").strip()
    password = data.get("password", "").strip()
    role = data.get("role", "user")
    display_name = data.get("display_name", "").strip()
    department = data.get("department", "General").strip()
    can_view_kb = data.get("can_view_kb", True)
    can_edit_kb = data.get("can_edit_kb", False)
    can_delete_kb = data.get("can_delete_kb", False)
    
    if not username or not password:
        return jsonify({"ok": False, "error": "กรุณากรอกชื่อผู้ใช้และรหัสผ่าน"}), 400
        
    if database.admin_create_user(username, password, role, display_name=display_name, can_view_kb=can_view_kb, can_edit_kb=can_edit_kb, can_delete_kb=can_delete_kb, department=department):
        database.log_event(f"Created user: {username} (Role: {role})", user=session.get("user"))
        return jsonify({"ok": True})
    return jsonify({"ok": False, "error": "ชื่อผู้ใช้นี้มีอยู่ในระบบแล้ว"}), 400

@app.route("/api/admin/users/<username>", methods=["PUT"])
@admin_required
def admin_update_user_route(username):
    data = request.json
    display_name = data.get("display_name")
    role = data.get("role")
    is_active = data.get("is_active")
    can_view_kb = data.get("can_view_kb")
    can_edit_kb = data.get("can_edit_kb")
    can_delete_kb = data.get("can_delete_kb")
    department = data.get("department")
    notes = data.get("notes")
    password = data.get("password")
    
    database.admin_update_user(username, display_name, role, is_active, notes, can_view_kb, can_edit_kb, can_delete_kb, department=department)
    if password:
        database.admin_reset_user_password(username, password)
        
    database.log_event(f"Updated user details for: {username}", user=session.get("user"))
    return jsonify({"ok": True})

@app.route("/api/admin/users/<username>/reset-password", methods=["POST"])
@admin_required
def admin_reset_password_route(username):
    data = request.json
    password = data.get("password")
    if not password:
        return jsonify({"ok": False, "error": "กรุณากรอกรหัสผ่านใหม่"}), 400
    
    database.admin_reset_user_password(username, password)
    database.log_event(f"Reset password for: {username}", user=session.get("user"))
    return jsonify({"ok": True})

@app.route("/api/admin/users/<username>", methods=["DELETE"])
@admin_required
def admin_delete_user_route(username):
    if username == session.get("user"):
        return jsonify({"ok": False, "error": "ไม่สามารถลบตัวเองได้"}), 400
        
    if database.admin_delete_user_complete(username):
        database.log_event(f"Deleted user: {username}", user=session.get("user"))
        return jsonify({"ok": True})
    return jsonify({"ok": False, "error": "ไม่สามารถลบผู้ใช้นี้ได้"}), 400


@app.route("/api/export")
def export_chat():
    try:
        history = database.get_history()
        if not history:
            return jsonify({"ok": False, "error": "ไม่มีประวัติการสนทนาสำหรับส่งออก"}), 400
            
        md_content = export_service.export_to_markdown(history)
        export_dir = Path("exports")
        file_path = export_service.save_export_file(md_content, export_dir)
        
        return send_from_directory(directory=str(export_dir), path=file_path.name, as_attachment=True)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/logs")
@admin_required
def get_logs():
    return jsonify({"logs": database.get_events()})

# --- Social Feed Routes ---
@app.route("/api/posts", methods=["GET", "POST"])
@login_required
def manage_posts():
    if request.method == "POST":
        # Handle multipart form data for file uploads
        content = request.form.get("content")
        author = session.get("user", request.form.get("author", "Anonymous"))  # Always use session user
        category = request.form.get("category", "General")
        link = request.form.get("link")
        
        if not content:
            return jsonify({"ok": False, "error": "Content required"}), 400
            
        attachments = []
        if 'files' in request.files:
            files = request.files.getlist('files')
            for file in files:
                if file.filename:
                    filename = secure_filename(file.filename)
                    # Add timestamp to avoid collisions
                    ts = int(time.time())
                    unique_filename = f"{ts}_{filename}"
                    save_path = os.path.join("uploads/social_feed", unique_filename)
                    file.save(save_path)
                    attachments.append({
                        "name": filename,
                        "path": f"/uploads/social_feed/{unique_filename}",
                        "type": file.content_type
                    })
            
        # Handle poll data
        poll_question = request.form.get("poll_question")
        poll_options_raw = request.form.get("poll_options")
        
        pid = database.add_post(content, author, category, link, attachments)
        
        if poll_question and poll_options_raw:
            try:
                import json
                poll_options = json.loads(poll_options_raw)
                if isinstance(poll_options, list) and len(poll_options) >= 2:
                    database.add_poll(pid, poll_question, poll_options)
                    database.log_event(f"Poll added to post: ID {pid}", user=author)
            except Exception as e:
                print(f"Error adding poll: {e}")

        database.log_event(f"New post created: ID {pid} by {author}", user=author)
        
        # --- New Post Broadcast ---
        profile = database.get_user_profile(author)
        display_name = profile.get("display_name", author)
        
        notification_db.notify_all_except(
            author, 
            'post', 
            'โพสต์ใหม่จาก ' + display_name, 
            f'หัวข้อ: {category} - "{content[:30]}..."',
            link='#feed'
        )
        # --- LINE Broadcast for Announcements & General Updates ---
        if category and category.lower() in ["announcement", "ประกาศ", "news", "ทั่วไป", "general", "it help"]:
            threading.Thread(target=broadcast_line_announcement, args=(
                "อัปเดตข่าวสารสำคัญ", 
                f"📢 มีประกาศใหม่ล่าสุด!\nหัวข้อ: {category}\nเนื้อหา: {content[:100]}...\nโดย: {display_name}"
            )).start()
        else:
            # Fallback broadcast for any other category
            threading.Thread(target=broadcast_line_announcement, args=(
                "โพสต์ใหม่ใน Feed", 
                f"✨ มีรายการใหม่น่าสนใจใน Feed!\nหมวดหมู่: {category}\nผู้โพสต์: {display_name}\nลองเข้าไปอ่านและแสดงความคิดเห็นได้ที่แอปนะครับ"
            )).start()
        
        # --- @Mentions in post ---
        import re
        mentions = re.findall(r'@(\w+)', content or '')
        mention_targets = [m for m in set(mentions) if m.lower() != author.lower()]
        if mention_targets:
            notification_db.notify_users(
                mention_targets,
                'mention',
                f'{display_name} แท็กคุณในโพสต์',
                f'"{(content or "")[:60]}..."',
                link='#feed'
            )
            # Push to LINE for mentions
            for target in mention_targets:
                threading.Thread(target=send_line_push_notification, args=(target, f'{display_name} แท็กคุณในโพสต์', f'"{(content or "")[:60]}..."')).start()
            batch_send_push_notification(mention_targets, f'{display_name} แท็กคุณในโพสต์', f'"{(content or "")[:40]}..."', url='#feed')
        
        return jsonify({"ok": True, "id": pid})
    
    cat = request.args.get("category", "All")
    posts = database.get_posts(cat)
    return jsonify({"posts": posts})

@app.route("/api/posts/<int:pid>", methods=["PUT", "DELETE"])
@login_required
def update_delete_post(pid):
    user = session.get("user")
    if request.method == "DELETE":
        ok = database.delete_post(pid, username=user, is_admin=is_admin())
        if ok:
            database.log_event(f"Post deleted: ID {pid}", user=user)
            return jsonify({"ok": True})
        return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์ลบโพสต์นี้"}), 403
    
    data = request.get_json(force=True)
    content = data.get("content")
    category = data.get("category")
    link = data.get("link")
    
    database.update_post(pid, content, category, link)
    database.log_event(f"Post updated: ID {pid}")
    return jsonify({"ok": True})

@app.route("/api/posts/<int:pid>/pin", methods=["POST"])
@admin_required
def pin_post_route(pid):
    database.toggle_pin(pid)
    return jsonify({"ok": True})

@app.route("/api/posts/<int:pid>/comments", methods=["GET", "POST"])
def manage_comments(pid):
    if request.method == "POST":
        data = request.get_json(force=True)
        content = data.get("content")
        author = session.get("user", data.get("author", "Anonymous"))  # Always use session user
        
        if not content:
            return jsonify({"ok": False, "error": "Comment content required"}), 400
            
        database.add_comment(pid, content, author)
        
        # --- Comment Notification ---
        posts = database.get_posts()
        post = next((p for p in posts if p["id"] == pid), None)
        if post and post["author"] != author:
            profile = database.get_user_profile(author)
            display_name = profile.get("display_name", author)
            notification_db.add_notification(
                post["author"],
                'comment',
                'มีคนแสดงความคิดเห็น',
                f'{display_name} แสดงความคิดเห็นในโพสต์ของคุณ: "{content[:30]}..."',
                link=f'#post-{pid}'
            )
            # Push to LINE
            threading.Thread(target=send_line_push_notification, args=(post["author"], 'มีคนแสดงความคิดเห็น', f'{display_name} แสดงความคิดเห็นในโพสต์: "{content[:30]}..."')).start()
            send_push_notification(post["author"], 'มีคนแสดงความคิดเห็น', f'{display_name} แสดงความคิดเห็นในโพสต์ของคุณ', url='#feed')
        
        # --- @Mentions in comment ---
        import re
        mentions = re.findall(r'@(\w+)', content or '')
        mention_targets = [m for m in set(mentions) if m.lower() != author.lower()]
        if mention_targets:
            notification_db.notify_users(
                mention_targets,
                'mention',
                f'{display_name} แท็กคุณในคอมเมนต์',
                f'"{(content or "")[:60]}..."',
                link=f'#post-{pid}'
            )
            batch_send_push_notification(mention_targets, f'{display_name} แท็กคุณในคอมเมนต์', f'"{(content or "")[:40]}..."', url=f'#feed')
            
            # --- LINE Push for Mentions in Comment ---
            for target in mention_targets:
                threading.Thread(target=send_line_push_notification, args=(target, f'{display_name} แท็กคุณในคอมเมนต์', f'"{(content or "")[:60]}..."')).start()
        
        return jsonify({"ok": True})
    
    return jsonify({"comments": database.get_comments(pid)})

@app.route("/api/posts/<int:pid>/view", methods=["POST"])
@login_required
def record_post_view_route(pid):
    user = session.get("user")
    database.record_post_view(pid, user)
    return jsonify({"ok": True})

@app.route("/api/posts/<int:pid>/views", methods=["GET"])
@login_required
def get_post_views_route(pid):
    views = database.get_post_views(pid)
    return jsonify({"ok": True, "views": views})

@app.route("/api/admin/analytics", methods=["GET"])
@admin_required
def get_admin_analytics():
    stats = database.get_analytics_data()
    return jsonify({"ok": True, "stats": stats})



@app.route("/api/posts/<int:pid>/comments/<int:cid>", methods=["DELETE"])
@login_required
def delete_comment_route(pid, cid):
    user = session.get("user")
    ok = database.delete_comment(cid, username=user, is_admin=is_admin())
    if ok:
        database.log_event(f"Comment deleted: ID {cid} from post {pid}", user=user)
        return jsonify({"ok": True})
    return jsonify({"ok": False, "error": "คุณไม่มีสิทธิ์ลบคอมเม้นนี้"}), 403

@app.route("/api/posts/<int:pid>/like", methods=["POST"])
def post_like(pid):
    user = session.get("user", "Current User")
    liked = database.toggle_like(pid, user)
    
    # --- Reaction Notification ---
    if liked:
        posts = database.get_posts()
        post = next((p for p in posts if p["id"] == pid), None)
        if post and post["author"] != user:
            notification_db.add_notification(
                post["author"],
                'like',
                'มีคนถูกใจโพสต์ของคุณ',
                f'{user} ถูกใจโพสต์ของคุณ',
                link=f'#post-{pid}'
            )
            # Push notification for like
            send_push_notification(post["author"], 'มีคนถูกใจโพสต์ของคุณ', f'{user} ถูกใจโพสต์ของคุณ', url='#feed')
            # --- LINE Push for Like ---
            threading.Thread(target=send_line_push_notification, args=(post["author"], 'มีคนถูกใจโพสต์ของคุณ', f'{user} ถูกใจโพสต์ของคุณ')).start()
    return jsonify({"ok": True, "liked": liked})

@app.route("/api/posts/<int:pid>/react", methods=["POST"])
@login_required
def post_react(pid):
    """Set/toggle an emoji reaction on a post."""
    user = session.get("user", "Anonymous")
    data = request.get_json(force=True)
    reaction = data.get("reaction", "like")
    
    reacted, reaction_type = database.set_reaction(pid, user, reaction)

    # Notification if reacted (not removed)
    if reacted:
        REACTION_LABELS = {
            'like': '👍 ถูกใจ', 'love': '❤️ รักเลย',
            'haha': '😂 ฮาเลย', 'wow': '😮 ทึ่ง',
            'sad': '😢 เศร้า', 'angry': '😡 โกรธ'
        }
        posts = database.get_posts()
        post = next((p for p in posts if p["id"] == pid), None)
        if post and post["author"] != user:
            notification_db.add_notification(
                post["author"],
                'like',
                f'มีคนแสดงความรู้สึกต่อโพสต์ของคุณ',
                f'{user} กด {REACTION_LABELS.get(reaction_type, reaction_type)} กับโพสต์ของคุณ',
                link=f'#post-{pid}'
            )
            # --- LINE Push for Reaction ---
            label = REACTION_LABELS.get(reaction_type, reaction_type)
            threading.Thread(target=send_line_push_notification, args=(post["author"], 'ความเคลื่อนไหวในโพสต์', f'{user} กด {label} กับโพสต์ของคุณ')).start()

    # Return updated counts grouped by reaction
    reactions = database.get_post_reactions(pid)
    counts = {}
    for r in reactions:
        rtype = r['reaction']
        counts[rtype] = counts.get(rtype, 0) + 1
    
    return jsonify({
        "ok": True, "reacted": reacted, "reaction": reaction_type,
        "counts": counts, "total": len(reactions)
    })

@app.route("/api/posts/<int:pid>/reactions", methods=["GET"])
@login_required
def get_reactions(pid):
    """Get all reactions for a post with user info."""
    user = session.get("user")
    reactions = database.get_post_reactions(pid)
    
    # Find current user's reaction
    my_reaction = None
    for r in reactions:
        if r['user'] == user:
            my_reaction = r['reaction']
            break
    
    # Group by reaction type
    counts = {}
    for r in reactions:
        rtype = r['reaction']
        counts[rtype] = counts.get(rtype, 0) + 1
    
    return jsonify({
        "ok": True, "reactions": reactions, "counts": counts,
        "total": len(reactions), "my_reaction": my_reaction
    })


@app.route("/api/posts/<int:pid>/summarize", methods=["POST"])
def summarize_post_route(pid):
    posts = database.get_posts()
    post = next((p for p in posts if p["id"] == pid), None)
    if not post:
        return jsonify({"ok": False, "error": "Post not found"}), 404
    
    content = post["content"]
    system_prompt = "ช่วยสรุปโพสต์นี้ให้เพื่อนๆ ในบริษัทอ่านเข้าใจง่ายๆ ภายใน 1-2 ประโยค ใช้ภาษาไทยที่เป็นธรรมชาติที่สุด"
    
    try:
        provider = ai_providers.get_provider()
        # Non-streaming call for summary
        full_summary = ""
        for chunk in provider.chat_stream(content, [], system_prompt):
            if chunk:
                full_summary += chunk
        
        database.update_post_summary(pid, full_summary)
        return jsonify({"ok": True, "summary": full_summary})
    except Exception as e:
        print(f"❌ Summarization Error: {e}")
        return jsonify({"ok": False, "error": str(e)}), 500


# --- Poll Routes ---
@app.route("/api/polls/<int:poll_id>/vote", methods=["POST"])
@login_required
def vote_poll_route(poll_id):
    data = request.get_json(force=True)
    option_id = data.get("option_id")
    user = session.get("user")
    
    if option_id is None:
        return jsonify({"ok": False, "error": "Missing option_id"}), 400
        
    # Get post_id for event logging
    ok = database.vote_poll(poll_id, option_id, user)
    if ok:
        database.log_event(f"User {user} voted in poll {poll_id}", user=user)
        return jsonify({"ok": True})
    return jsonify({"ok": False, "error": "คุณลงคะแนนไม่สำเร็จ"}), 500

@app.route("/api/polls/<int:poll_id>/user_vote")
@login_required
def get_user_vote_route(poll_id):
    user = session.get("user")
    vote = database.get_user_vote(poll_id, user)
    return jsonify({"option_id": vote})


@app.route("/api/feed/summarize", methods=["POST"])
@login_required
def feed_daily_summary():
    user = session.get("user", "Admin")
    data = database.get_daily_activities()
    
    posts = data.get("posts", [])
    schedules = data.get("schedules", [])
    
    if not posts and not schedules:
        return jsonify({"ok": True, "summary": "วันนี้ยังไม่มีความเคลื่อนไหวใหม่ในองค์กรครับ"})
    
    # Prepare prompt
    posts_text = "\n".join([f"- {p['author']} โพสต์ใน {p['category']}: {p['content'][:100]}" for p in posts])
    schedules_text = "\n".join([f"- {s['title']} ({s['category']}) วันที่ {s['date']} เวลา {s['time']}" for s in schedules])
    
    prompt = f"""คุณคือเลขาส่วนตัวอัจฉริยะ ทำหน้าที่สรุป 'Morning Brief' หรือภาพรวมกิจกรรมล่าสุดใน Feed ให้พนักงาน
    
    ข้อมูลโพสต์ในฟีด 24 ชม. ที่ผ่านมา:
    {posts_text}
    
    ตารางนัดหมาย/กิจกรรมที่กำลังเกิดขึ้น:
    {schedules_text}
    
    สรุปให้เป็นกันเองเหมือนเล่าข่าวในที่ทำงาน (ไม่เกิน 4-5 ประโยค) เน้นประเด็นที่ทุกคนควรรู้เพื่อให้ทำงานได้ราบรื่นในวันนี้ ใช้ภาษาไทยที่สละสลวย
    """
    
    try:
        provider = ai_providers.get_provider()
        full_summary = ""
        # Using a default system prompt for character consistency
        system_prompt = "คุณคือ 'AI-Assistant' ผู้ช่วยสรุปข่าวสารและกิจกรรมภายในองค์กรที่รอบรู้และสุภาพ"
        
        for chunk in provider.chat_stream(prompt, [], system_prompt):
            if chunk:
                full_summary += chunk
        
        return jsonify({"ok": True, "summary": full_summary})
    except Exception as e:
        print(f"❌ Daily Summary Error: {e}")
        return jsonify({"ok": False, "error": str(e)}), 500






# --- Self-Service Profile Update ---
@app.route("/api/profile", methods=["GET"])
@login_required
def get_my_profile():
    username = session.get("user")
    profile = database.get_user_profile(username)
    return jsonify({"ok": True, "profile": profile})

@app.route("/api/profile/update", methods=["POST"])
@login_required
def update_my_profile():
    username = session.get("user")
    
    # Handle multipart form (avatar file upload) or JSON
    display_name = None
    avatar_url = None
    department = None
    
    if request.content_type and 'multipart/form-data' in request.content_type:
        display_name = request.form.get("display_name")
        department = request.form.get("department")
        
        # Handle avatar file upload
        avatar_file = request.files.get("avatar")
        if avatar_file and avatar_file.filename:
            filename = secure_filename(avatar_file.filename)
            ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else 'png'
            avatar_filename = f"avatar_{username}_{int(time.time())}.{ext}"
            upload_dir = os.path.join("static", "uploads", "avatars")
            os.makedirs(upload_dir, exist_ok=True)
            avatar_path = os.path.join(upload_dir, avatar_filename)
            avatar_file.save(avatar_path)
            avatar_url = f"/static/uploads/avatars/{avatar_filename}"
    else:
        data = request.get_json(force=True)
        display_name = data.get("display_name")
        avatar_url = data.get("avatar_url") 
        department = data.get("department")
    
    database.update_user_profile(username, display_name=display_name, avatar_url=avatar_url, department=department)
    profile = database.get_user_profile(username)
    database.log_event(f"Profile updated by {username}", username)
    return jsonify({"ok": True, "profile": profile})

# --- Last Seen / User Status ---
@app.route("/api/users/status")
@login_required
def get_all_users_status():
    """Get online/offline status + last_seen for all users."""
    try:
        conn = sqlite3.connect("chat_history.db")
        cursor = conn.cursor()
        cursor.execute("""
            SELECT us.username, us.is_online, us.last_activity, up.display_name, up.avatar_url
            FROM user_status us
            LEFT JOIN user_profiles up ON us.username = up.username COLLATE NOCASE
            ORDER BY us.is_online DESC, us.last_activity DESC
        """)
        rows = cursor.fetchall()
        conn.close()
        
        users = []
        for row in rows:
            users.append({
                "username": row[0],
                "is_online": bool(row[1]),
                "last_activity": row[2],
                "display_name": row[3] or row[0],
                "avatar_url": row[4]
            })
        return jsonify({"ok": True, "users": users})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# --- Unified Chat Routes ---
@app.route("/api/notifications/subscribe", methods=["POST"])
@login_required
def subscribe_push():
    data = request.get_json(force=True)
    subscription = data.get("subscription")
    if not subscription:
        return jsonify({"ok": False, "error": "Subscription required"}), 400
    
    username = session.get("user")
    database.add_push_subscription(username, json.dumps(subscription))
    return jsonify({"ok": True})

@app.route("/api/chat/list")
@login_required
def get_chat_list():
    user = session.get("user")
    data = database.get_rooms_for_user(user)
    return jsonify({"ok": True, "rooms": data["rooms"], "contacts": data["contacts"]})

@app.route("/api/chat/users")
@login_required
def get_chat_users():
    # Return all users except current one
    current_user_raw = session.get("user", "")
    current_user_lower = current_user_raw.lower()
    
    # 1. Collect all potential usernames from hardcoded USERS and database
    all_usernames = set()
    
    # Add hardcoded users
    for u in USERS.keys():
        all_usernames.add(u.capitalize())
        
    all_usernames.add("AI-Assistant")
    
    # Add database users
    try:
        db_users = database.get_all_usernames()
        for du in db_users:
            all_usernames.add(du)
    except Exception as e:
        print(f"Error fetching db usernames: {e}")
        
    profiles = []
    # 2. Fetch profiles for each unique user except current one
    for u in all_usernames:
        if u.lower() != current_user_lower:
            try:
                prof = database.get_user_profile(u)
                if prof:
                    profiles.append(prof)
            except Exception as e:
                print(f"Error fetching profile for {u}: {e}")
                # Fallback profile
                profiles.append({
                    "username": u,
                    "display_name": u.capitalize(),
                    "avatar_url": None,
                    "department": "General"
                })
                
    # Sort for better UI
    profiles.sort(key=lambda x: (x.get("display_name") or x.get("username", "")).lower())
    
    return jsonify({"ok": True, "users": profiles})


@app.route("/api/groups/create", methods=["POST"])
@login_required
def create_new_group():
    user = session.get("user")
    data = request.get_json(force=True)
    name = data.get("name", "Unnamed Group")
    members = data.get("members", []) # List of usernames
    
    room_id = database.create_room(name, user, members)
    database.log_event(f"Created group: {name} (ID: {room_id})", user=user)
    return jsonify({"ok": True, "room_id": room_id})


# --- Kanban (Task Status) ---
@app.route("/api/schedules/<int:sched_id>/status", methods=["PATCH"])
@login_required
def patch_schedule_status(sched_id):
    data = request.json or {}
    status = data.get("status")
    if not status:
        return jsonify({"ok": False, "error": "Status required"}), 400
    
    user = session.get("user")
    if database.update_schedule_status(sched_id, status, user):
        return jsonify({"ok": True})
    return jsonify({"ok": False, "error": "Could not update status"}), 403

# --- Internal Knowledge Wiki ---
@app.route("/api/wiki", methods=["GET"])
@login_required
def list_wiki_pages():
    pages = database.get_wiki_pages()
    return jsonify({"pages": pages})

@app.route("/api/wiki", methods=["POST"])
@login_required
def create_wiki_page_route():
    data = request.json or {}
    title = data.get("title", "").strip()
    content = data.get("content", "").strip()
    category_id = data.get("category_id")
    
    if not title or not content:
        return jsonify({"ok": False, "error": "Title and content required"}), 400
        
    user = session.get("user")
    page_id, slug = database.create_wiki_page(title, content, user, category_id)
    
    # Ingest into RAG engine
    rag_engine.ingest_text(content, source_name=f"Wiki: {title}", category_id=category_id)
    
    database.log_event(f"Created Wiki page: {title}", user=user)
    return jsonify({"ok": True, "slug": slug})

@app.route("/api/wiki/<slug>", methods=["GET"])
@login_required
def get_wiki_page_route(slug):
    page = database.get_wiki_page(slug)
    if not page:
        return jsonify({"ok": False, "error": "Page not found"}), 404
    return jsonify({"ok": True, "page": page})

@app.route("/api/wiki/<slug>", methods=["PUT"])
@login_required
def update_wiki_page_route(slug):
    data = request.json or {}
    title = data.get("title", "").strip()
    content = data.get("content", "").strip()
    category_id = data.get("category_id")
    
    if not title or not content:
        return jsonify({"ok": False, "error": "Title and content required"}), 400
        
    if database.update_wiki_page(slug, title, content, category_id):
        # Update RAG engine
        rag_engine.ingest_text(content, source_name=f"Wiki: {title}", category_id=category_id)
        
        database.log_event(f"Updated Wiki page: {title}", user=session.get("user"))
        return jsonify({"ok": True})
    return jsonify({"ok": False, "error": "Update failed"}), 404

@app.route("/api/wiki/<slug>", methods=["DELETE"])
@login_required
def delete_wiki_page_route(slug):
    page = database.get_wiki_page(slug)
    if not page:
        return jsonify({"ok": False}), 404
        
    if database.delete_wiki_page(slug):
        # Remove from RAG
        rag_engine._kb.delete_by_source(f"Wiki: {page['title']}")
        database.log_event(f"Deleted Wiki page: {slug}", user=session.get("user"))
        return jsonify({"ok": True})
    return jsonify({"ok": False}), 500


@app.route("/api/groups/<int:room_id>/members", methods=["POST"])
@login_required
def add_members_to_group(room_id):
    user = session.get("user")
    data = request.get_json(force=True)
    members = data.get("members", [])
    
    if not members:
        return jsonify({"ok": False, "error": "ไม่ได้ระบุสมาชิกที่ต้องการเพิ่ม"}), 400
        
    database.add_room_members(room_id, members)
    database.log_event(f"Added {len(members)} members to room {room_id}", user=user)
    return jsonify({"ok": True})

@app.route("/api/groups/<int:gid>/profile", methods=["POST"])
@login_required
def update_group_profile_route(gid):
    user = session.get("user")
    # Check if owner
    room = next((r for r in database.get_rooms_for_user(user)["rooms"] if r["id"] == gid), None)
    if not room or room["owner"] != user and user != "Admin":
        return jsonify({"ok": False, "error": "Unauthorized"}), 403
        
    name = request.form.get("name")
    avatar_url = room.get("avatar_url") # Default to current
    if 'avatar' in request.files:
        file = request.files['avatar']
        if file.filename:
            group_dir = Path("uploads/group_profiles")
            group_dir.mkdir(parents=True, exist_ok=True)
            filename = secure_filename(f"group_{gid}_{int(time.time())}_{file.filename}")
            file.save(str(group_dir / filename))
            avatar_url = f"/uploads/group_profiles/{filename}"
            
    database.update_group_profile(gid, name=name, avatar_url=avatar_url)
    return jsonify({"ok": True, "avatar_url": avatar_url})

@app.route("/api/groups/<int:gid>", methods=["DELETE"])
@admin_required
def delete_group_route(gid):
    user = session.get("user")
    ok = database.delete_room(gid)
    if ok:
        database.log_event(f"Deleted group room ID: {gid}", user=user)
        return jsonify({"ok": True})
    return jsonify({"ok": False, "error": "ไม่พบกลุ่มที่ต้องการลบ"}), 404

@app.route("/api/chat/rooms/<int:room_id>/members")
@login_required
def get_room_members(room_id):
    """Returns members of a room with their profile info for @mention autocomplete."""
    user = session.get("user")
    # Verify user is a member of this room
    rooms_data = database.get_rooms_for_user(user)
    room = next((r for r in rooms_data["rooms"] if r["id"] == room_id), None)
    if not room:
        return jsonify({"ok": False, "error": "Not a member"}), 403
    
    conn = sqlite3.connect(database.DB_PATH)
    cursor = conn.cursor()
    cursor.execute("""
        SELECT rm.username, COALESCE(up.display_name, rm.username) as display_name, up.avatar_url
        FROM room_members rm
        LEFT JOIN user_profiles up ON LOWER(up.username) = LOWER(rm.username)
        WHERE rm.room_id = ?
        ORDER BY display_name ASC
    """, (room_id,))
    rows = cursor.fetchall()
    conn.close()
    
    members = [{"username": r[0], "display_name": r[1], "avatar_url": r[2]} for r in rows]
    return jsonify({"ok": True, "users": members})

@app.route("/api/groups/<int:room_id>/members/<username>", methods=["DELETE"])
@admin_required
def remove_room_member_route(room_id, username):
    ok, msg = database.remove_room_member(room_id, username)
    if ok:
        database.log_event(f"Removed user {username} from room {room_id}", user=session.get("user"))
        return jsonify({"ok": True})
    return jsonify({"ok": False, "error": msg}), 400

@app.route("/api/chat/messages/<ctype>/<cid>")
@login_required
def get_chat_messages_route(ctype, cid):
    user = session.get("user")
    peek = request.args.get("peek") == "1"
    
    last_reads = []
    if ctype == "room":
        # Check membership
        rooms = database.get_rooms_for_user(user)["rooms"]
        if not any(r["id"] == int(cid) for r in rooms):
            return jsonify({"ok": False, "error": "Not a member"}), 403
        messages = database.get_room_messages(int(cid))
        last_reads = database.get_room_reader_avatars(int(cid))
        if not peek:
            database.mark_room_read(int(cid), user, messages[-1]["id"] if messages else 0)
    else: # dm
        messages = database.get_dm_history(user, cid)
        if not peek:
            database.mark_dm_read(user, cid, messages[-1]["id"] if messages else 0)
            
    return jsonify({"ok": True, "messages": messages, "last_reads": last_reads})

@app.route("/api/chat/send", methods=["POST"])
@login_required
def send_unified_message():
    user = session.get("user")
    ctype = request.form.get("type") # 'room' or 'dm'
    cid = request.form.get("id") # roomId (int) or recipient (str)
    text = request.form.get("message", "").strip()
    reply_to_id = request.form.get("reply_to_id")
    if reply_to_id:
        try: reply_to_id = int(reply_to_id)
        except: reply_to_id = None
    files = request.files.getlist("files")
    
    attachments = []
    if files:
        upload_subdir = "uploads/group_chat" if ctype == "room" else "uploads/dm_chat"
        Path(upload_subdir).mkdir(parents=True, exist_ok=True)
        for file in files:
            if file and file.filename:
                filename = secure_filename(f"{user}_{int(time.time())}_{file.filename}")
                file.save(os.path.join(upload_subdir, filename))
                if file.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp')):
                    ftype = "image"
                elif file.filename.lower().endswith(('.mp3', '.wav', '.ogg', '.m4a', '.webm')):
                    ftype = "audio"
                else:
                    ftype = "file"
                attachments.append({
                    "url": f"/{upload_subdir}/{filename}",
                    "name": file.filename,
                    "type": ftype
                })

    if not text and not attachments:
        return jsonify({"ok": False, "error": "Message cannot be empty"}), 400

    if ctype == "room":
        mid = database.add_room_message(int(cid), user, text, attachments, reply_to_id=reply_to_id)
        
        # Offload mention detection and notifications to a thread
        def handle_room_notifications(m_id, u, t, r_id):
            if not t: return
            import re
            mentions = [m.lower() for m in re.findall(r'@(\w+)', t)]
            all_users = database.admin_get_all_users()
            
            profile = database.get_user_profile(u)
            display_name = profile.get('display_name', u)
            room_data = next((r for r in database.get_rooms_for_user(u)['rooms'] if r['id'] == r_id), None)
            room_name = room_data['name'] if room_data else 'ห้องกลุ่ม'
            
            target_usernames = []
            for target_user in all_users:
                target_username = target_user['username'].lower()
                target_display = (target_user['display_name'] or "").lower()
                if target_username == u.lower(): continue
                
                if target_username in mentions or (target_display and f"@{target_display}" in t.lower()):
                    target_usernames.append(target_user['username'])

            if target_usernames:
                notification_db.notify_users(
                    target_usernames, 'mention', f'ถูก Mention ในห้อง {room_name}',
                    f'{display_name} กล่าวถึงคุณ: "{t[:40]}"',
                    link='#chat'
                )
                batch_send_push_notification(target_usernames, f'Mention ใน {room_name}', f'{display_name}: {t[:40]}', url='#chat')
                
                # --- LINE Push for Mentions in Group ---
                for target in target_usernames:
                    threading.Thread(target=send_line_push_notification, args=(target, f'Mention ใน {room_name}', f'{display_name}: {t[:40]}')).start()

        threading.Thread(target=handle_room_notifications, args=(mid, user, text, int(cid))).start()

        # AI Bot response
        if text and '@bot' in text.lower():
            # Find the first image in attachments to send to AI
            ai_image = None
            ai_mime = "image/jpeg"
            if attachments:
                img_att = next((a for a in attachments if a.get("type") == "image"), None)
                if img_att:
                    try:
                        rel_path = img_att["url"].lstrip("/")
                        abs_path = os.path.join(os.getcwd(), rel_path)
                        if os.path.exists(abs_path):
                            with open(abs_path, "rb") as f:
                                ai_image = f.read()
                            if rel_path.lower().endswith(".png"): ai_mime = "image/png"
                            elif rel_path.lower().endswith(".webp"): ai_mime = "image/webp"
                            elif rel_path.lower().endswith(".gif"): ai_mime = "image/gif"
                    except Exception as e:
                        print(f"⚠️ Error loading image for AI: {e}")

            threading.Thread(target=handle_bot_response, args=(int(cid), text, "room", user, ai_image, ai_mime)).start()
            
    else: # dm
        mid = database.save_dm(user, cid, text, attachments, reply_to_id=reply_to_id)
        
        def handle_dm_notifications(u, target_id, t):
            if target_id.lower() == 'ai-assistant': return
            
            profile = database.get_user_profile(u)
            display_name = profile.get('display_name', u)
            preview = (t[:40] + '...') if (t and len(t) > 40) else (t or '(ไฟล์แนบ)')
            
            # Check for mention in DM
            target_profile = database.get_user_profile(target_id)
            target_user_lower = target_id.lower()
            target_display_lower = (target_profile.get('display_name') or "").lower()
            
            if t and (f"@{target_user_lower}" in t.lower() or (target_display_lower and f"@{target_display_lower}" in t.lower())):
                notification_db.add_notification(
                    target_id, 'mention', f'ถูก Mention จาก {display_name}',
                    preview, link='#chat'
                )
            
            # For DMs, we ALWAYS send a push notification
            send_push_notification(target_id, f'ข้อความใหม่จาก {display_name}', preview, url='#chat')
            
            # --- LINE Push for DMs ---
            threading.Thread(target=send_line_push_notification, args=(target_id, f'ข้อความใหม่จาก {display_name}', preview)).start()

        threading.Thread(target=handle_dm_notifications, args=(user, cid, text)).start()
        
        # AI Bot response in DM
        if cid.lower() == 'ai-assistant' or (text and '@bot' in text.lower()):
            # Find the first image in attachments to send to AI
            ai_image = None
            ai_mime = "image/jpeg"
            if attachments:
                img_att = next((a for a in attachments if a.get("type") == "image"), None)
                if img_att:
                    try:
                        # Extract absolute path from relative URL
                        rel_path = img_att["url"].lstrip("/")
                        abs_path = os.path.join(os.getcwd(), rel_path)
                        if os.path.exists(abs_path):
                            with open(abs_path, "rb") as f:
                                ai_image = f.read()
                            # Guess mime type
                            if rel_path.lower().endswith(".png"): ai_mime = "image/png"
                            elif rel_path.lower().endswith(".webp"): ai_mime = "image/webp"
                            elif rel_path.lower().endswith(".gif"): ai_mime = "image/gif"
                    except Exception as e:
                        print(f"⚠️ Error loading image for AI: {e}")

            threading.Thread(target=handle_bot_response, args=(cid, text, "dm", user, ai_image, ai_mime)).start()
            
    # Emit socket signal for real-time update
    socket_room = f"room_{cid}" if ctype == "room" else f"dm_{user}_{cid}"
    
    # Common response data
    res_data = {
        "id": mid,
        "sender": user,
        "username": user,
        "text": text,
        "attachments": attachments,
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "reply_to_id": reply_to_id
    }
    
    # Add reply context if exists
    if reply_to_id:
        # We need to fetch reply context info for real-time update
        # For simplicity we fetch it from the message data if we don't have it handy
        # But usually, it's safer to have the DB return it or just send what we have
        # In this context, let's just send the ID, and the frontend will already have the message in its local state.
        # Actually, for correctness, we should include sender and text for the recipient who might not have it.
        try:
            conn = sqlite3.connect(database.DB_PATH)
            cur = conn.cursor()
            table = "room_messages" if ctype == "room" else "private_messages"
            sender_col = "username" if ctype == "room" else "sender"
            cur.execute(f"SELECT {sender_col}, text FROM {table} WHERE id = ?", (reply_to_id,))
            row = cur.fetchone()
            if row:
                res_data["reply_sender"] = row[0]
                res_data["reply_text"] = row[1]
            conn.close()
        except: pass

    if ctype == 'room':
        res_data["room_id"] = int(cid)
        res_data["type"] = "room"
        socketio.emit('new_message', res_data, room=socket_room)
    else:
        # To Recipient
        res_dm_data = res_data.copy()
        res_dm_data["type"] = "dm"
        res_dm_data["sender_id"] = user
        socketio.emit('new_message', res_dm_data, room=f"dm_{cid}")
        # To Sender (self)
        socketio.emit('new_message', res_dm_data, room=f"dm_{user}")

    return jsonify({"ok": True, "message_id": mid})

@app.route("/api/chat/read", methods=["POST"])
@login_required
def chat_mark_read():
    data = request.json or {}
    ctype = data.get("type") # 'room' or 'dm'
    cid = data.get("id") # roomId or recipient
    max_id = data.get("max_id") # highest message id seen
    
    if not ctype or not cid or not max_id:
        return jsonify({"ok": False, "error": "Missing parameters"}), 400
        
    username = session.get("user")
    if ctype == "room":
        database.mark_room_read(int(cid), username, int(max_id))
        socketio.emit("read_update", {"type": "room", "id": int(cid), "user": username, "max_id": int(max_id)}, room=f"room_{cid}")
    else:
        database.mark_dm_read(username, cid, int(max_id))
        # For DM, we notify the sender that we read their messages
        socketio.emit("read_update", {"type": "dm", "id": username, "user": username, "max_id": int(max_id)}, room=f"dm_{cid}")
        
    return jsonify({"ok": True})

def log_bot(message):
    try:
        with open("bot_debug.log", "a", encoding="utf-8") as f:
            f.write(f"[{get_current_time()}] {message}\n")
    except:
        pass

def handle_bot_response(cid, user_text, ctype, original_sender=None, image_data=None, mime_type="image/jpeg"):
    """Generates an AI response for the chat, now with Vision support."""
    try:
        log_bot(f"🤖 Bot generating response for: {user_text[:50]}... User: {original_sender}")
        # 1. Retrieve context with permission filter
        rag_filter = get_rag_filter(original_sender)
        context, sources = rag_engine.retrieve_context(user_text, where=rag_filter)
        log_bot(f"🔍 RAG context retrieved ({len(context)} chars). Sources: {len(sources)}")
        
        # Greeting Detection
        is_greeting = any(x in user_text.lower() for x in ["สวัสดี", "hi", "hello", "หวัดดี"]) and len(user_text) < 15

        now = get_current_time() # Format: 2026-03-19 18:35:22
        system_prompt = (
            f"คุณคือ 'AI-Assistant' (OrgChat AI) ผู้ช่วยประจำองค์กรที่มีความเฉลียวฉลาด รอบรู้ และมีความเป็นมนุษย์สูง\n"
            f"เวลาปัจจุบันคือ: {now}\n\n"
            "กฎเหล็กในการสื่อสาร (สำคัญมาก):\n"
            "1. หากผู้ใช้แค่ 'ทักทาย' (เช่น สวัสดี, Hi) ให้ตอบทักทายกลับสั้นๆ อย่างเป็นธรรมชาติและสุภาพ โดย 'ห้าม' แสดงข้อมูลปฏิทินหรือข้อมูลเชิงลึกในเอกสารหากไม่ได้ถูกถาม\n"
            "2. ใช้ภาษาไทยระดับเจ้าของภาษา (Native Thai) สไตล์คนทำงานที่เก่งและใจดี หลีกเลี่ยงสำนวนแปลตรงตัว\n"
            "3. ห้ามใช้ 'ผม/ดิฉัน', 'ฉันสามารถช่วยคุณได้อย่างไร', หรือคำว่า 'มีอะไรหรือยัง?' ในการทักทาย\n"
            "4. แทนตัวเองว่า 'ผม' และแทนผู้ใช้ว่า 'คุณ' ลงท้ายด้วย 'ครับ' ทุกครั้ง\n"
            "5. ใช้ข้อมูลใน Context และปฏิทิน 'เฉพาะเมื่อจำเป็นและเกี่ยวข้อง' เท่านั้น\n"
            "6. หากไม่พบข้อมูลที่ต้องการ ให้บอกตรงๆ ว่า 'ไม่พบข้อมูลในระบบครับ' พร้อมอาสาจะช่วยเรื่องอื่นแทน\n\n"
        )
        
        # Inject Schedules into context only if NOT a simple greeting or if specifically asked
        current_date_str = get_current_time().split()[0]
        calendar_info = f"=== วันสำคัญและกิจกรรมองค์กร (วันนี้คือ: {current_date_str}) ===\n"
        for date, name in THAI_HOLIDAYS_2026.items():
            calendar_info += f"- {date}: {name} (วันหยุดนักขัตฤกษ์)\n"

        current_user = original_sender or "System"
        schedules = database.get_schedules(current_user)
        if schedules:
            for s in schedules[-15:]:
                calendar_info += f"- {s['date']} {s['time']}: {s['title']} ({s['desc']})\n"
        
        calendar_info += "=== สิ้นสุดข้อมูลปฏิทิน ===\n\n"
        
        if not is_greeting:
            system_prompt += calendar_info
            system_prompt += f"Context:\n{context}"
        else:
            # If greeting, still provide basic context but tell AI to lead the conversation
            system_prompt += "หมายเหตุ: นี่คือการทักทายเบื้องต้น ให้ตอบกลับอย่างละมุนและสุภาพ พร้อมถามว่ามีอะไรให้ช่วยไหม\n"
            # Optional: provide only today's events if it's a greeting to be helpful
            today_sched = [s for s in schedules if s['date'] == current_date_str]
            if today_sched:
                system_prompt += f"กิจกรรมวันนี้ของคุณ: " + ", ".join([s['title'] for s in today_sched]) + "\n"
        
        # 2. Get history for context (optional but better)
        history = []
        if ctype == "room":
            raw_history = database.get_room_messages(cid)[-5:]
            for m in raw_history:
                role = "model" if m["username"] == "AI-Assistant" else "user"
                history.append({"role": role, "text": m["text"]})
        else:
            raw_history = database.get_dm_history(original_sender or "System", cid)[-5:]
            for m in raw_history:
                role = "model" if m["username"] == "AI-Assistant" else "user"
                history.append({"role": role, "text": m["text"]})

        # 3. Call AI
        log_bot(f"📡 Calling AI provider ({os.environ.get('AI_PROVIDER', 'gemini')}) (Vision: {image_data is not None})...")
        provider = ai_providers.get_provider()
        
        # If image is present, adjust prompt for Vision Expert mode
        if image_data:
            system_prompt += (
                "\n\n[Vision Mode Enabled]\n"
                "ผู้ใช้ได้ส่งรูปภาพมาให้คุณวิเคราะห์ โปรดอธิบายสิ่งที่เห็นในรูปตามความเหมาะสม "
                "หากเป็นเอกสารหรือบิล ให้สรุปข้อมูลสำคัญ ตัวเลข หรือรายการออกมาให้ชัดเจนที่สุด "
                "หากมีสิ่งของหรือสถานที่ ให้อธิบายลักษณะเด่นของสิ่งนั้นๆ ครับ"
            )
        
        response_text = ""
        
        # Mark as typing
        socketio.emit('bot_typing', {'room_id': cid, 'type': ctype}, room=f"room_{cid}" if ctype == "room" else f"dm_{cid}")
        
        # Use simple generator if image is present (Vision calls are usually atomic, but our provider supports streaming)
        # We pass image_data down to the provider
        stream = provider.chat_stream(user_text, history, system_prompt, image_data=image_data, mime_type=mime_type)
        bot_key = (cid, ctype)
        if bot_key not in _typing_status: _typing_status[bot_key] = {}
        _typing_status[bot_key]["AI-Assistant"] = time.time()
        
        for chunk in provider.chat_stream(user_text, history, system_prompt):
            response_text += chunk
            # Emit chunk to client for real-time streaming
            socket_room = f"room_{cid}" if ctype == "room" else f"dm_{original_sender}"
            socketio.emit('ai_chunk', {
                "cid": cid,
                "ctype": ctype,
                "chunk": chunk,
                "is_start": (len(response_text) == len(chunk))
            }, room=socket_room)
            
            # Keep typing status alive
            _typing_status[bot_key]["AI-Assistant"] = time.time()
            
        if not response_text:
            log_bot("⚠️ Provider returned empty response.")
            response_text = "ขออภัยครับ ผมไม่พบข้อมูลที่เกี่ยวข้องในระบบฐานความรู้ขององค์กร"
            socketio.emit('ai_chunk', {"cid": cid, "ctype": ctype, "chunk": response_text, "is_start": True}, room=socket_room)
        else:
            log_bot(f"✅ Bot response generated ({len(response_text)} chars).")
            # Clear bot typing status and notify end of stream
            socketio.emit('ai_done', {"cid": cid, "ctype": ctype}, room=socket_room)
            if "AI-Assistant" in _typing_status.get(bot_key, {}):
                _typing_status[bot_key]["AI-Assistant"] = 0 

        # 4. Save response
        if ctype == "room":
            database.add_room_message(cid, "AI-Assistant", response_text, [])
        else:
            database.save_dm("AI-Assistant", original_sender, response_text, [])
            # NEW: Send push notification to the user who mentioned the bot in DM
            if original_sender:
                send_push_notification(original_sender, "AI-Assistant ตอบกลับคุณ", response_text[:100], url='#chat')
            
    except Exception as e:
        import traceback
        log_bot(f"❌ Bot Response Error: {e}")
        log_bot(traceback.format_exc())
        error_msg = f"ขออภัยครับ เกิดข้อผิดพลาดทางเทคนิค: {str(e)}"
        if ctype == "room":
            database.add_room_message(cid, "AI-Assistant", error_msg, [])
        else:
            database.save_dm("AI-Assistant", original_sender, error_msg, [])

@app.route("/api/chat/unread")
@login_required
def get_unread_counts_route():
    user = session.get("user")
    return jsonify({"ok": True, "unread": database.get_unread_counts(user)})

# --- Typing Indicator (in-memory, expires in 5s) ---
_typing_status = {}  # {(chat_id, chat_type): {username: timestamp}}

# --- Typing Indicator (Moving to Socket.IO strictly) ---
@socketio.on('typing')
def handle_typing_socket(data):
    user = session.get("user")
    if not user: return
    cid = data.get("id")
    ctype = data.get("type", "room")
    room = f"room_{cid}" if ctype == "room" else f"dm_{user}"
    
    # Broadcast to others in the same chat
    emit('is_typing', {"username": user, "id": cid, "type": ctype}, room=room, include_self=False)

# 🎨 Whiteboard Real-time Sync
@socketio.on('join_whiteboard')
def on_join_whiteboard():
    join_room('whiteboard_global')

@socketio.on('draw')
def handle_draw(data):
    # data: {x0, y0, x1, y1, color, size}
    emit('draw_update', data, room='whiteboard_global', include_self=False)

@socketio.on('whiteboard_cursor')
def handle_wb_cursor(data):
    if "user" in session:
        data["username"] = session["user"]
        emit('wb_cursor_update', data, room='whiteboard_global', include_self=False)

@socketio.on('clear_whiteboard')
def handle_clear_whiteboard():
    emit('whiteboard_cleared', room='whiteboard_global', include_self=False)

# Keep simple GET for initial sync if needed, but remove the polling-heavy check if desired
@app.route("/api/chat/typing", methods=["GET"])
@login_required
def get_typing():
    return jsonify({"typing": []}) # Polling disabled in favor of sockets

# --- Calendar Reminder (check schedules due in 15 min) ---
@app.route("/api/schedules/reminders")
@login_required
def check_reminders():
    from datetime import datetime, timedelta
    user = session.get("user")
    schedules = database.get_schedules(user)
    now = datetime.now()
    reminders = []
    for s in schedules:
        try:
            dt_str = f"{s['date']} {s.get('time', '00:00')}"
            dt = datetime.strptime(dt_str, "%Y-%m-%d %H:%M")
            diff = (dt - now).total_seconds() / 60
            if 0 < diff <= 15:
                reminders.append({
                    "id": s["id"],
                    "title": s["title"],
                    "minutes_left": int(diff),
                    "time": s.get("time", "")
                })
        except Exception:
            pass
    return jsonify({"reminders": reminders})


@app.route("/api/chat", methods=["POST"])
@login_required
def chat():
    provider = os.environ.get("AI_PROVIDER", "gemini").lower()
    api_key = _get_gemini_api_key() if provider == "gemini" else True
    print(f"💬 Chat request received. Provider: {provider}. API Key present: {bool(api_key)}")
    if not api_key:
        return jsonify({"ok": False, "error": "ยังไม่ได้ตั้งค่า Gemini API Key"}), 400

    data = request.get_json(force=True)
    import time
    t0 = time.time() # Start Total Timer
    
    question = (data.get("message") or "").strip()
    history = data.get("history", []) # list of {role, text}
    persona_id = data.get("persona_id")
    image_data_raw = data.get("image_data") # Base64 string
    mime_type = data.get("mime_type", "image/jpeg")

    # Handle Base64 image data if present
    image_bytes = None
    if image_data_raw:
        try:
            import base64
            # Strip data:image/xxx;base64, prefix if exists
            if "," in image_data_raw:
                header, image_data_raw = image_data_raw.split(",", 1)
                # Try to guess mime type from header if not provided
                if "mime_type" not in data and "image/" in header:
                    mime_type = header.split(":")[1].split(";")[0]
            
            image_bytes = base64.b64decode(image_data_raw)
            print(f"🖼️ Image received: {len(image_bytes)} bytes. Mime: {mime_type}")
        except Exception as e:
            print(f"⚠️ Error decoding image: {e}")

    if not question and not image_bytes:
        return jsonify({"ok": False, "error": "กรุณาพิมพ์คำถามหรือส่งรูปภาพ"}), 400

    current_user = session.get("user", "Admin")

    # 1. Retrieve Context from KB with permission filter
    import re
    rag_filter = get_rag_filter(current_user)
    
    file_match = re.search(r"ช่วยสรุปเนื้อหาและดึงข้อมูลสำคัญจากไฟล์ ID:\s*([a-f0-9\-]+)", question)
    if file_match:
        file_id = file_match.group(1).strip()
        print(f"📄 Detected File Summary Request for ID: {file_id}")
        
        file_filter = {"file_id": file_id}
        if rag_filter:
            file_filter = {"$and": [rag_filter, {"file_id": file_id}]}
            
        chunks = rag_engine.get_all_chunks(limit=50, where=file_filter)
        
        context_parts = []
        sources = []
        seen_s = set()
        
        for r in chunks:
            loc = r.get("location", "Document")
            dept = r.get("department", "General")
            src = r.get("source", "?")
            fid = r.get("file_id", str(file_id))
            
            context_parts.append(f"--- แหล่งที่มา: {src} [{loc}] (File ID: {fid}) ---\n{r.get('text', '')}")
            
            s_key = f"{fid}_{loc}"
            if s_key not in seen_s:
                sources.append({"file_id": str(fid), "name": src, "location": loc, "department": dept, "type": r.get("type", "file")})
                seen_s.add(s_key)
                
        context = "\n\n".join(context_parts)
        if sources:
            file_name = sources[0]["name"]
            question += f"\n(ระบบแนบข้อมูล 50 ส่วนแรกจากไฟล์ '{file_name}' มาให้พิจารณาแล้ว ใน context โปรดสรุปภาพรวมให้ที)"
    else:
        # Skip RAG context if it's just an image question without much text

        # 1. Retrieve Context
        context, sources = "", []
        common_greets = ["สวัสดี", "หวัดดี", "ว่าไง", "ทักทาย", "hi", "hello", "hey"]
        is_greeting = any(g in (question or "").lower() for g in common_greets)
        
        if question and len(question) > 5 and not is_greeting:
            import sys
            rag_filter = get_rag_filter(current_user)
            context, sources = rag_engine.retrieve_context(question, where=rag_filter)
            print(f"⏱️ DEBUG: RAG Retrieval took {time.time() - t0:.3f}s", flush=True)
            sys.stdout.flush()
        else:
            context, sources = "", []
            if is_greeting: 
                print("⚡ Fast Path: Skipping RAG", flush=True)
                import sys; sys.stdout.flush()

        t1 = time.time()

    # Build base system prompt
    now_str = get_current_time()
    
    # Check for Persona
    persona_prompt = ""
    persona_name = "AI Assistant"
    if persona_id:
        persona = database.get_persona(persona_id)
        if persona:
            persona_prompt = persona.get("system_prompt", "")
            persona_name = persona.get("name", "AI Assistant")

    agentic_info = (
        f"\n[สำคัญ] วันนี้คือ {now_str}. หากผู้ใช้พูดถึงวัน (เช่น พรุ่งนี้, วันจันทร์หน้า) ให้คุณคำนวณวันที่จริง (YYYY-MM-DD) ออกมาให้ถูกต้อง\n"
        "หากผู้ใช้ต้องการ 'นัดหมาย' 'จองคิว' 'แจ้งเตือน' หรือ 'สร้างงาน' ให้คุณสรุปงานนั้นเป็น JSON ท้ายคำตอบโดยใช้รูปแบบดังนี้:\n"
        '[CALENDAR_ACTION]{"title": "...", "date": "YYYY-MM-DD", "time": "HH:MM", "desc": "..."}[/CALENDAR_ACTION]\n'
        "คุณสามารถตอบคำถามทั่วไปพร้อมกับสร้างนัดหมายไปพร้อมกันได้เลย\n"
    )

    # 2. Build final prompt with Persona, Weather, and Schedules
    weather_ctx = get_weather_context()
    schedules = database.get_schedules(current_user)

    # --- CACHED FILENAME INVENTORY (Speed Fix) ---
    global _kb_filenames_cache
    if '_kb_filenames_cache' not in globals() or not _kb_filenames_cache:
        try:
            inv_data = rag_engine._kb.collection.get(include=['metadatas'], limit=100)
            _kb_filenames_cache = sorted(list(set([m.get("source") for m in inv_data.get("metadatas", []) if m.get("source")])))
        except:
            _kb_filenames_cache = []
    
    kb_inventory = f"\n[คลังข้อมูล]: คุณมีสิทธิ์เข้าถึงไฟล์: {', '.join(_kb_filenames_cache)}\n" if _kb_filenames_cache else ""
    # ---------------------------------------------

    system_prompt = (
        f"คุณคือผู้ช่วยอัจฉริยะ {persona_name} ประจำองค์กร\n"
        f"วันนี้คือ {now_str}. "
        f"{agentic_info}\n"
        f"{kb_inventory}\n"
    )

    if weather_ctx:
        system_prompt += (
            "\n[ข้อมูลสภาพอากาศล่าสุด (Real-time)]\n"
            f"{weather_ctx}\n"
            "คำแนะนำ: หากผู้ใช้ถามเรื่องอากาศ ฝนตก หรืออุณหภูมิ ให้ใช้ข้อมูลจริงด้านบนนี้ตอบทันที "
            "โดยทำหน้าที่เสมือนคุณมีเซนเซอร์อากาศติดตัว ไม่ต้องระบุว่าเป็นข้อมูลจากคลังความรู้\n\n"
        )

    if persona_prompt:
        system_prompt += f"{persona_prompt}\n"
    else:
        system_prompt += (
            "ตอบคำถามจากข้อมูลบริษัท กิจกรรม และปฏิทินที่ได้รับอย่างแม่นยำและสุภาพ\n"
            "ตอบเป็นภาษาไทย สรุปใจความสำคัญ ไม่เยิ่นเย้อ และเน้นความถูกต้อง\n"
        )

    if image_bytes:
        system_prompt += (
            "\n\n[Vision Mode Enabled]\n"
            "ผู้ใช้ได้ส่งรูปภาพมาให้คุณวิเคราะห์ โปรดอธิบายสิ่งที่เห็นในรูปตามความเหมาะสม "
            "หากเป็นเอกสารหรือบิล ให้สรุปข้อมูลสำคัญ ตัวเลข หรือรายการออกมาให้ชัดเจนที่สุด "
            "หากมีสิ่งของหรือสถานที่ ให้อธิบายลักษณะเด่นของสิ่งนั้นๆ ครับ"
        )

    # Add Schedules
    system_prompt += "=== วันสำคัญและกิจกรรมองค์กร ===\n"
    for date, name in THAI_HOLIDAYS_2026.items():
        system_prompt += f"- {date}: {name} (วันสำคัญ/วันหยุดนักขัตฤกษ์)\n"
    if schedules:
        for s in schedules[-10:]:
            system_prompt += f"- {s['date']} {s['time']}: {s['title']} ({s['desc']})\n"
    system_prompt += "=== สิ้นสุดข้อมูลปฏิทิน ===\n\n"

    if context:
        system_prompt += (
            "=== ข้อมูลอ้างอิงจากเอกสารองค์กร ===\n"
            + context
            + "\n=== สิ้นสุดข้อมูลอ้างอิง ===\n"
        )
    else:
        system_prompt += "\n(ขณะนี้ยังไม่มีเอกสารอ้างอิงที่เกี่ยวข้องในระบบ)\n"

    try:
        provider_obj = ai_providers.get_provider()
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 400

    print(f"🤖 Chat session started for {current_user} with provider: {os.environ.get('AI_PROVIDER', 'gemini')} (Vision: {image_bytes is not None})")
    
    def generate():
        try:
            # First, send the sources
            print(f"🟢 Yielding sources: {sources}")
            yield f"data: {json.dumps({'sources': sources})}\n\n"
            
            # Save user message
            print(f"⏱️ DEBUG: Starting AI Call. Elapsed from start: {time.time() - t0:.3f}s", flush=True); import sys; sys.stdout.flush()
            u_sid = database.save_message("user", question, username=current_user)
            yield f"data: {json.dumps({'user_id': u_sid})}\n\n"

            # Stream chunks directly to client
            print(f"⏳ Calling {os.environ.get('AI_PROVIDER', 'gemini')} provider (Real-time Streaming)...", flush=True); sys.stdout.flush()
            ai_call_start = time.time()
            response_stream = provider_obj.chat_stream(question, history, system_prompt, image_data=image_bytes, mime_type=mime_type)
            bot_full_text = ""
            
            got_first_chunk = False
            for chunk in response_stream:
                if chunk:
                    if not got_first_chunk:
                        print(f"✅ DEBUG: First Chunk from AI received in {time.time() - ai_call_start:.3f}s (Total from msg start: {time.time() - t0:.3f}s)", flush=True)
                        sys.stdout.flush()
                        got_first_chunk = True
                    bot_full_text += chunk
                    # Send chunk immediately
                    yield f"data: {json.dumps({'content': chunk})}\n\n"
            
            print(f"🏁 Stream finished. Total length: {len(bot_full_text)}")
            
            # Save bot message to DB
            if bot_full_text:
                b_sid = database.save_message("bot", bot_full_text, sources=sources, username=current_user)
                # Final signal to frontend
                yield f"data: {json.dumps({'done': True, 'bot_id': b_sid})}\n\n"
            
        except Exception as e:
            error_msg = str(e)
            if "429" in error_msg or "quota" in error_msg.lower():
                error_msg = "ขออภัยครับ ตอนนี้โควต้าการใช้งาน Gemini API (Free Tier) ของคุณเต็มแล้ว กรุณารอสักครู่ (ประมาณ 1 นาที) หรือเช็คการตั้งค่า API Key ครับ"
            print(f"❌ Error in chat stream: {e}")
            yield f"data: {json.dumps({'error': error_msg})}\n\n"

    return app.response_class(generate(), mimetype='text/event-stream')

@app.route("/api/chat/typing", methods=["POST"])
@login_required
def chat_typing():
    # Dummy route to stop 405 errors and improve UI feel
    return jsonify({"ok": True})

@app.route("/api/chat/rooms/pin/<int:mid>", methods=["POST"])
@login_required
def pin_room_message(mid):
    # For now, let's allow anyone to pin, or we could restrict to room owners
    database.toggle_room_message_pin(mid)
    return jsonify({"ok": True})

@app.route("/api/chat/pin/<ctype>/<int:mid>", methods=["POST"])
@login_required
def pin_chat_message(ctype, mid):
    if ctype == 'room':
        database.toggle_room_message_pin(mid)
    else: # dm
        database.toggle_private_message_pin(mid)
    return jsonify({"ok": True})

@app.route("/api/chat/delete/<ctype>/<int:mid>", methods=["DELETE", "POST"])
@login_required
def delete_chat_message(ctype, mid):
    user = session.get("user")
    if ctype == 'ai':
        ok = database.delete_message(mid, username=user)
    elif ctype == 'room':
        ok = database.delete_room_message(mid, username=user, is_admin=is_admin())
    elif ctype == 'dm':
        ok = database.delete_private_message(mid, username=user, is_admin=is_admin())
    else:
        # Fallback for old calls if any still exist without ctype (though technically this route wouldn't match)
        ok = database.delete_message(mid, username=user)
        
    if ok:
        return jsonify({"ok": True})
    return jsonify({"ok": False, "error": "ไม่พบข้อความหรือคุณไม่มีสิทธิ์ในการลบ"}), 404

# ─── Message Edit API ─────────────────────────────────────
@app.route("/api/chat/edit/<ctype>/<int:mid>", methods=["PUT", "POST"])
@login_required
def edit_chat_message(ctype, mid):
    """Edit a chat message (room or dm). Only the owner or admin can edit."""
    user = session.get("user")
    data = request.get_json(force=True) or {}
    new_text = (data.get("text") or "").strip()
    if not new_text:
        return jsonify({"ok": False, "error": "ข้อความไม่สามารถว่างได้"}), 400

    if ctype == 'room':
        ok = database.edit_room_message(mid, new_text, user, is_admin_user=is_admin())
    elif ctype == 'dm':
        ok = database.edit_private_message(mid, new_text, user, is_admin_user=is_admin())
    else:
        return jsonify({"ok": False, "error": "ประเภทห้องไม่รองรับ"}), 400

    if ok:
        database.log_event(f"Message {mid} edited by {user}", user=user)
        return jsonify({"ok": True, "text": new_text})
    return jsonify({"ok": False, "error": "ไม่พบข้อความหรือคุณไม่มีสิทธิ์แก้ไข"}), 403


# ─── Global Search API ────────────────────────────────────
@app.route("/api/search/global")
@login_required
def global_search_route():
    """Search across posts, schedules, chat, and DMs."""
    q = request.args.get("q", "").strip()
    user = session.get("user", "Anonymous")
    
    # Log search for insights
    if q:
        database.log_search(q, user)
        
    if len(q) < 2:
        return jsonify({"ok": True, "results": []})
        
    try:
        # Text DB search
        db_results = database.global_search(q, user, limit=20)
        # KB search (if user has access)
        kb_results = []
        if can_view_knowledge_base():
            try:
                rag_filter = get_rag_filter(user)
                raw = rag_engine.search_kb(q, n_results=5, where=rag_filter)
                for r in raw:
                    kb_results.append({
                        "type": "kb",
                        "id": r.get("file_id", ""),
                        "text": r.get("text", "")[:200],
                        "author": "Knowledge Base",
                        "category": r.get("name", r.get("source", "")),
                        "timestamp": "",
                        "link": "#kb"
                    })
            except Exception:
                pass
        all_results = db_results + kb_results
        return jsonify({"ok": True, "results": all_results, "query": q})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route("/api/admin/search-insights")
@admin_required
def search_insights_route():
    insights = database.get_search_insights()
    return jsonify({"ok": True, "insights": insights})

@app.route("/api/ai/extract-tasks", methods=["POST"])
@login_required
def extract_tasks_route():
    data = request.get_json()
    chat_text = data.get("text", "")
    if not chat_text:
        return jsonify({"ok": False, "error": "No text provided"}), 400
        
    prompt = (
        "วิเคราะห์ข้อความแชทต่อไปนี้ และสกัด 'Action Items' หรือ 'งานที่ต้องทำ' ออกมา "
        "หากพบงาน ให้สรุปเป็น JSON array ของวัตถุที่มีฟิลด์ 'title' และ 'description' (ภาษาไทย) "
        "หากไม่พบงาน ให้ส่งเป็น array ว่าง [] "
        "ตอบเฉพาะ JSON เท่านั้น:\n\n"
        f"ข้อความ: {chat_text}"
    )
    
    try:
        provider = ai_providers.get_provider()
        full_response = ""
        for chunk in provider.chat_stream(prompt, [], "คุณคือผู้ช่วยสกัดงานจากบทสนทนา"):
            full_response += chunk
            
        # Clean up JSON from markdown if exists
        clean_json = full_response.strip()
        if "```json" in clean_json:
            clean_json = clean_json.split("```json")[1].split("```")[0].strip()
        elif "```" in clean_json:
            clean_json = clean_json.split("```")[1].split("```")[0].strip()
            
        tasks = json.loads(clean_json)
        return jsonify({"ok": True, "tasks": tasks})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/notifications", methods=["GET"])
@login_required
def get_notifications_route():
    user = session.get("user")
    notifs = notification_db.get_notifications(user)
    unread_count = sum(1 for n in notifs if not n.get('is_read'))
    return jsonify({"notifications": notifs, "unread_count": unread_count})

@app.route("/api/notifications/<int:notif_id>/read", methods=["POST"])
@login_required
def mark_notif_read_route(notif_id):
    notification_db.mark_notification_read(notif_id)
    return jsonify({"ok": True})

@app.route("/api/notifications/read_all", methods=["POST"])
@login_required
def mark_all_read_route():
    user = session.get("user")
    notification_db.mark_all_notifications_read(user)
    return jsonify({"ok": True})

@app.route("/api/notifications/<int:notif_id>", methods=["DELETE"])
@login_required
def delete_notif_route(notif_id):
    user = session.get("user")
    notification_db.delete_notification(notif_id, user)
    return jsonify({"ok": True})

@app.route("/api/notifications/delete_all", methods=["DELETE"])
@login_required
def delete_all_notifs_route():
    user = session.get("user")
    notification_db.delete_all_notifications(user)
    return jsonify({"ok": True})

# ─── Leave Request System ──────────────────────────────────
@app.route("/api/leave/request", methods=["POST"])
@login_required
def api_leave_request():
    user = session.get("user")
    data = request.get_json()
    l_type = data.get("type", "Sick")
    s_date = data.get("start_date")
    e_date = data.get("end_date")
    reason = data.get("reason", "")
    
    if not s_date or not e_date:
        return jsonify({"ok": False, "error": "กรุณากรอกวันที่ให้ครบถ้วน"}), 400
        
    leave_id = database.create_leave_request(user, l_type, s_date, e_date, reason)
    if leave_id:
        # Notify admins
        admins = database.get_all_admins()
        for admin in admins:
            if admin.lower() == user.lower(): continue
            notification_db.add_notification(
                admin, 
                "leave_request", 
                "คำขอลาใหม่", 
                f"คุณ {user} ได้ส่งคำขอลา {l_type} ตั้งแต่ {s_date} ถึง {e_date}", 
                "/#admin"
            )
            # send_push_notification(admin, "คำขอลาใหม่", f"คุณ {user} ได้ส่งคำขอลา {l_type}", "/#admin")
        return jsonify({"ok": True, "id": leave_id})
    return jsonify({"ok": False, "error": "เกิดข้อผิดพลาดในการส่งคำขอ"}), 500

@app.route("/api/leave/my")
@login_required
def api_leave_my():
    user = session.get("user")
    leaves = database.get_user_leaves(user)
    return jsonify({"ok": True, "leaves": leaves})

@app.route("/api/admin/leave/all")
@admin_required
def api_admin_leave_all():
    leaves = database.get_all_leaves()
    return jsonify({"ok": True, "leaves": leaves})

@app.route("/api/admin/leave/status", methods=["POST"])
@admin_required
def api_admin_leave_status():
    admin_user = session.get("user")
    data = request.get_json()
    leave_id = data.get("id")
    status = data.get("status") # approved/rejected
    note = data.get("note", "")
    
    if not leave_id or status not in ["approved", "rejected"]:
        return jsonify({"ok": False, "error": "ข้อมูลไม่ถูกต้อง"}), 400
        
    ok = database.update_leave_status(leave_id, status, admin_user, note)
    if ok:
        leave = database.get_leave_request(leave_id)
        if leave:
            notification_db.add_notification(
                leave["username"],
                "leave_status",
                f"คำขอลาของคุณได้รับการ{ 'อนุมัติ' if status == 'approved' else 'ปฏิเสธ' }",
                f"คำขอลาวันที่ {leave['start_date']} ถึง {leave['end_date']} ได้รับการ{ 'อนุมัติ' if status == 'approved' else 'ปฏิเสธ' } โดย {admin_user}",
                "/#leave"
            )
        return jsonify({"ok": True})
    return jsonify({"ok": False, "error": "ไม่พบคำขอลาที่ระบุ"}), 404

@app.route("/api/leave/comment", methods=["POST"])
@login_required
def api_leave_add_comment():
    user = session.get("user")
    data = request.get_json()
    leave_id = data.get("leave_id")
    comment = data.get("comment", "").strip()
    
    if not leave_id or not comment:
        return jsonify({"ok": False, "error": "ข้อมูลไม่ครบถ้วน"}), 400
        
    database.add_leave_comment(leave_id, user, comment)
    
    # Notify other party
    leave = database.get_leave_request(leave_id)
    if leave:
        target = ""
        msg = f"มีข้อความใหม่ในคำขอลา: {comment[:50]}..."
        if user.lower() == leave["username"].lower():
            # User commented -> Notify admins
            admins = database.get_all_admins()
            for admin in admins:
                notification_db.add_notification(admin, "leave_comment", "ข้อความใหม่จากพนักงาน", msg, "/#admin")
        else:
            # Admin commented -> Notify user
            notification_db.add_notification(leave["username"], "leave_comment", "ข้อความใหม่จากผู้ดูแล", msg, "/#leave")
            
    return jsonify({"ok": True})

@app.route("/api/leave/comments/<int:leave_id>")
@login_required
def api_leave_get_comments(leave_id):
    comments = database.get_leave_comments(leave_id)
    return jsonify({"ok": True, "comments": comments})

# ─── Lunch Randomizer API ────────────────────────────────
@app.route("/api/lunch/random")
@login_required
def api_lunch_random():
    return jsonify(database.get_random_lunch() or {"name": "ยังไม่มีร้านอาหารในระบบ"})

@app.route("/api/lunch/all")
@login_required
def api_lunch_all():
    return jsonify(database.get_all_lunch_places())

@app.route("/api/lunch/add", methods=["POST"])
@login_required
def api_lunch_add():
    data = request.get_json() or {}
    name = data.get("name")
    if not name: return jsonify({"ok": False}), 400
    database.add_lunch_place(name, data.get("type", ""), data.get("location", ""), session.get("user"))
    return jsonify({"ok": True})

@app.route("/api/admin/notifications/clear_all", methods=["POST"])
@admin_required
def admin_clear_all_notifications_route():
    count = notification_db.admin_clear_all_notifications()
    database.log_event(f"Admin cleared ALL notifications ({count} items)", user=session.get("user"))
    return jsonify({"status": "success", "count": count})


# ─── Admin Chat Management API ───────────────────────────────

@app.route("/api/admin/chat/overview", methods=["GET"])
@admin_required
def admin_chat_overview():
    rooms = database.admin_get_all_rooms()
    dm_pairs = database.admin_get_all_dm_pairs()
    return jsonify({"ok": True, "rooms": rooms, "dm_pairs": dm_pairs})

# --- AI Chat ---
@app.route("/api/admin/chat/ai", methods=["GET"])
@admin_required
def admin_ai_messages():
    username = request.args.get("username")
    msgs = database.admin_get_all_ai_messages(limit=200)
    if username:
        msgs = [m for m in msgs if m["username"] == username]
    return jsonify({"ok": True, "messages": msgs})

@app.route("/api/admin/chat/ai/delete/<int:mid>", methods=["DELETE", "POST"])
@admin_required
def admin_delete_ai_msg(mid):
    ok = database.delete_message(mid, username="Admin")
    return jsonify({"ok": ok})

@app.route("/api/admin/chat/ai/clear", methods=["POST"])
@admin_required
def admin_clear_ai():
    data = request.get_json(silent=True) or {}
    count = database.admin_clear_ai_chat(username=data.get("username"))
    return jsonify({"ok": True, "deleted": count})

# --- Room Chat ---
@app.route("/api/admin/chat/room/<int:room_id>", methods=["GET"])
@admin_required
def admin_room_messages(room_id):
    msgs = database.admin_get_room_messages(room_id)
    return jsonify({"ok": True, "messages": msgs})

@app.route("/api/admin/chat/room/delete/<int:mid>", methods=["DELETE", "POST"])
@admin_required
def admin_delete_room_msg(mid):
    ok = database.admin_delete_room_message(mid)
    return jsonify({"ok": ok})

@app.route("/api/admin/chat/room/<int:room_id>/clear", methods=["POST"])
@admin_required
def admin_clear_room(room_id):
    count = database.admin_clear_room_messages(room_id)
    return jsonify({"ok": True, "deleted": count})

# --- DM Chat ---
@app.route("/api/admin/chat/dm", methods=["GET"])
@admin_required
def admin_dm_messages():
    u1 = request.args.get("user1", "")
    u2 = request.args.get("user2", "")
    if not u1 or not u2:
        return jsonify({"ok": False, "error": "user1 and user2 required"}), 400
    msgs = database.admin_get_dm_messages(u1, u2)
    return jsonify({"ok": True, "messages": msgs})

@app.route("/api/admin/chat/dm/delete/<int:mid>", methods=["DELETE", "POST"])
@admin_required
def admin_delete_dm_msg(mid):
    ok = database.admin_delete_dm_message(mid)
    return jsonify({"ok": ok})

@app.route("/api/admin/chat/dm/clear", methods=["POST"])
@admin_required
def admin_clear_dm_route():
    data = request.get_json(silent=True) or {}
    u1 = data.get("user1", "")
    u2 = data.get("user2", "")
    if not u1 or not u2:
        return jsonify({"ok": False, "error": "user1 and user2 required"}), 400
    count = database.admin_clear_dm(u1, u2)
    return jsonify({"ok": True, "deleted": count})








@app.route("/api/admin/users/<username>/toggle-active", methods=["POST"])
@admin_required
def admin_toggle_active_route(username):
    if username == "Admin":
        return jsonify({"ok": False, "error": "ไม่สามารถปิดบัญชี Admin ได้"}), 400
    settings = database.get_user_setting(username)
    new_active = not settings.get("is_active", True)
    database.admin_update_user(username, is_active=new_active)
    admin_user = session.get("user", "Admin")
    status_th = "เปิด" if new_active else "ปิด"
    database.log_event(f"Admin {status_th}การใช้งาน user: {username}", user=admin_user)
    return jsonify({"ok": True, "is_active": new_active})



@app.route("/manifest.json")
def serve_manifest():
    return send_from_directory("static", "manifest.json")


# ─── Socket.IO Events ──────────────────────────
@socketio.on('join')
def on_join(data):
    room = data.get('room')
    if room:
        join_room(room)
        print(f"👤 User {session.get('user')} joined room: {room}")

@socketio.on('leave')
def on_leave(data):
    room = data.get('room')
    if room:
        leave_room(room)

@socketio.on('chat_message')
def handle_chat_message(data):
    # This is a real-time signal, actual DB save is still via REST for reliability
    # but we can broadcast it here to avoid polling lag
    room = data.get('room')
    if room:
        emit('new_message', data, room=room)

# ─────────────── WebSocket Events for Real-Time Features ───────────────

@socketio.on('user_online')
def handle_user_online(data):
    """Handle user coming online."""
    try:
        username = session.get("user")
        if not username:
            return emit('error', {'msg': 'Not logged in'})
        
        room = data.get('room', 'General')
        database.set_user_online(username, room)
        
        # User joins their own personal room for WebRTC signaling
        join_room(f"user_{username}")
        
        # Notify all users in the room about new Online status
        online_users = database.get_online_users()
        emit('user_status_updated', {
            'username': username,
            'status': 'online',
            'online_users': online_users
        }, room=room)
        
        print(f"[ONLINE] {username} joined room: {room}")
    except Exception as e:
        print(f"[ERROR] user_online: {e}")
        emit('error', {'msg': str(e)})

@socketio.on('user_offline')
def handle_user_offline():
    """Handle user going offline."""
    try:
        username = session.get("user")
        if not username:
            return
        
        database.set_user_offline(username)
        online_users = database.get_online_users()
        
        # Notify all connected clients about offline status
        emit('user_status_updated', {
            'username': username,
            'status': 'offline',
            'online_users': online_users
        }, broadcast=True)
        
        print(f"[OFFLINE] {username} went offline")
    except Exception as e:
        print(f"[ERROR] user_offline: {e}")

@socketio.on('user_typing')
def handle_user_typing(data):
    """Handle typing indicator."""
    try:
        username = session.get("user")
        if not username:
            return
        
        room_id = data.get('room_id', 1)
        database.set_user_typing(username, room_id)
        
        # Broadcast typing indicator to room
        emit('user_typing', {
            'username': username,
            'room_id': room_id,
            'display_name': data.get('display_name', username)
        }, room=f"room_{room_id}")
        
    except Exception as e:
        print(f"[ERROR] user_typing: {e}")

@socketio.on('user_stopped_typing')
def handle_user_stopped_typing(data):
    """Clear typing indicator."""
    try:
        username = session.get("user")
        if not username:
            return
        
        room_id = data.get('room_id', 1)
        database.clear_user_typing(username)
        
        # Broadcast to room
        emit('user_stopped_typing', {
            'username': username,
            'room_id': room_id
        }, room=f"room_{room_id}")
        
    except Exception as e:
        print(f"[ERROR] user_stopped_typing: {e}")

@socketio.on('message_read')
def handle_message_read(data):
    """Mark message as read (read receipt)."""
    try:
        username = session.get("user")
        if not username:
            return
        
        message_id = data.get('message_id')
        room_id = data.get('room_id')
        
        if message_id:
            database.mark_message_as_read(message_id, username, room_id)
            
            # Get read receipts for this message
            receipts = database.get_message_read_receipts(message_id)
            
            # Broadcast read receipts to room
            emit('message_read_receipt', {
                'message_id': message_id,
                'username': username,
                'read_by_count': len(receipts),
                'read_receipts': receipts
            }, room=f"room_{room_id}" if room_id else None)
            
            print(f"[READ] {username} read message {message_id}")
        
    except Exception as e:
        print(f"[ERROR] message_read: {e}")

@socketio.on('message_read_dm')
def handle_private_message_read(data):
    """Mark private message as read."""
    try:
        username = session.get("user")
        if not username:
            return
        
        message_id = data.get('message_id')
        if message_id:
            database.mark_private_message_as_read(message_id)
            
            # Emit read receipt to sender
            sender = data.get('sender')
            emit('dm_read_receipt', {
                'message_id': message_id,
                'read_by': username,
                'read_at': time.strftime("%Y-%m-%d %H:%M:%S")
            }, room=f"dm_{sender}_{username}" if sender else None)
            
            print(f"[DM_READ] {username} read DM {message_id}")
            
    except Exception as e:
        print(f"[ERROR] message_read_dm: {e}")

@socketio.on('join_room')
def on_join_room(data):
    """Join a room group for broadcasts."""
    try:
        room = data.get('room', '1')
        username = session.get("user")
        join_room(f"room_{room}")
        print(f"[JOIN] {username} joined room group: {room}")
        emit('message', {'text': f'{username} joined'}, room=f"room_{room}")
    except Exception as e:
        print(f"[ERROR] join_room: {e}")

@socketio.on('get_online_users')
def handle_get_online_users():
    """Get current list of online users."""
    try:
        online_users = database.get_online_users()
        emit('online_users_list', {
            'users': online_users,
            'count': len(online_users),
            'timestamp': time.strftime("%Y-%m-%d %H:%M:%S")
        })
    except Exception as e:
        print(f"[ERROR] get_online_users: {e}")
        emit('error', {'msg': str(e)})

# ─── WebRTC Signaling ────────────────────────
@socketio.on('webrtc_signal')
def handle_webrtc_signal(data):
    """
    Signaling server for WebRTC Voice/Video Calls.
    """
    # Prefer sender from payload, fallback to session
    sender = data.get('sender') or session.get("user")
    if not sender:
        print(f"⚠️ [WebRTC] Signal dropped: No sender identified. Data={data}")
        return
        
    target = data.get('target')
    if not target:
        print(f"⚠️ [WebRTC] Signal dropped: No target for signal {data.get('type')} from {sender}")
        return
        
    target_room = f"user_{target}"
    data['sender'] = sender  # Ensure sender is attached
    
    print(f"📡 [WebRTC] {data.get('type').upper()} from {sender} -> {target}")
    
    # If it's a new call (offer), trigger a push notification to wake up the recipient
    if data.get('type') == 'offer':
        try:
            # Call the local function defined at the top of app.py
            call_type_th = "วิดีโอคอล" if data.get('callType') == 'video' else "สายโทรเข้า"
            send_push_notification(
                target, 
                f"📞 {sender}", 
                f"กำลังมี{call_type_th}จาก {sender}...", 
                tag="incoming_call"
            )
        except Exception as e:
            print(f"⚠️ [WebRTC] Push notification failed: {e}")

    emit('webrtc_signal', data, room=target_room)

@app.route('/sw.js')
def serve_sw():
    return send_from_directory('static', 'sw.js', mimetype='application/javascript')


# ─────────────────────── Kanban API ───────────────────────

database.kanban_init_db()

@app.route("/api/kanban/board", methods=["GET"])
@login_required
def kanban_get_board():
    return jsonify({"ok": True, "columns": database.kanban_get_board()})

@app.route("/api/kanban/columns", methods=["POST"])
@login_required
def kanban_add_column():
    data = request.json or {}
    title = data.get("title", "").strip()
    color = data.get("color", "#6366f1")
    if not title:
        return jsonify({"ok": False, "error": "กรุณาระบุชื่อคอลัมน์"}), 400
    col_id = database.kanban_add_column(title, color, created_by=session.get("user", "System"))
    return jsonify({"ok": True, "id": col_id})

@app.route("/api/kanban/columns/<int:col_id>", methods=["PUT"])
@login_required
def kanban_update_column(col_id):
    data = request.json or {}
    database.kanban_update_column(col_id, title=data.get("title"), color=data.get("color"))
    return jsonify({"ok": True})

@app.route("/api/kanban/columns/<int:col_id>", methods=["DELETE"])
@login_required
def kanban_delete_column(col_id):
    ok = database.kanban_delete_column(col_id)
    return jsonify({"ok": ok})

@app.route("/api/kanban/columns/reorder", methods=["POST"])
@login_required
def kanban_reorder_columns():
    data = request.json or {}
    order = data.get("order", [])
    database.kanban_reorder_columns(order)
    return jsonify({"ok": True})

@app.route("/api/kanban/cards", methods=["POST"])
@login_required
def kanban_add_card():
    data = request.json or {}
    column_id = data.get("column_id")
    title = data.get("title", "").strip()
    if not column_id or not title:
        return jsonify({"ok": False, "error": "Missing column_id or title"}), 400
    
    current_user = session.get("user", "System")
    assignees_raw = data.get("assignee", "")
    if assignees_raw and isinstance(assignees_raw, str):
        assignee_list = [a.strip() for a in assignees_raw.split(',') if a.strip()]
    else:
        assignee_list = []
        
    card_id = database.kanban_add_card(
        column_id=int(column_id),
        title=title,
        description=data.get("description", ""),
        priority=data.get("priority", "medium"),
        assignee=",".join(assignee_list) if assignee_list else "",
        due_date=data.get("due_date", ""),
        labels=data.get("labels", ""),
        color=data.get("color", ""),
        created_by=current_user,
        is_done=int(data.get("is_done", 0))
    )
    
    if assignee_list:
        for a in assignee_list:
            if a != current_user:
                try:
                    notification_db.add_notification(
                        a,
                        "kanban",
                        "มอบหมายงานใหม่ / Kanban",
                        f"{current_user} ได้มอบหมายงาน '{title}' ให้กับคุณ",
                        link="#view=kanban"
                    )
                    batch_send_push_notification([a], "มอบหมายงานใหม่ / Kanban", f"{current_user}: {title}", url="#view=kanban")
                    # --- Send to LINE ---
                    threading.Thread(target=send_line_push_notification, args=(
                        a, 
                        "งานใหม่ (Kanban)", 
                        f"{current_user} ได้มอบหมายงาน '{title}' ให้กับคุณ"
                    ), kwargs={
                        "fields": {
                            "งาน": title,
                            "ผู้มอบหมาย": current_user,
                            "สถานะ": "รอดำเนินการ (To Do)"
                        }
                    }).start()
                except Exception as e:
                    print(f"Error sending kanban notification to {a}: {e}")
            
    return jsonify({"ok": True, "id": card_id})

@app.route("/api/kanban/cards/<int:card_id>", methods=["PUT"])
@login_required
def kanban_update_card(card_id):
    data = request.json or {}
    
    old_card = database.kanban_get_card(card_id)
    if not old_card:
        return jsonify({"ok": False, "error": "ไม่พบบทความ"}), 404
        
    database.kanban_update_card(card_id, **data)
    
    current_user = session.get("user", "System")
    assignees_raw = data.get("assignee", "")
    if isinstance(assignees_raw, str):
        new_assignee_list = [a.strip() for a in assignees_raw.split(',') if a.strip()]
    else:
        new_assignee_list = []
        
    old_assignee_list = (old_card.get("assignee") or "").split(',')
    old_assignee_list = [a.strip() for a in old_assignee_list if a.strip()]
    
    # Notify ONLY newly added assignees
    added_assignees = [a for a in new_assignee_list if a not in old_assignee_list and a != current_user]
    
    for a in added_assignees:
        try:
            notification_db.add_notification(
                a,
                "kanban",
                "มอบหมายงาน / Kanban",
                f"{current_user} ได้มอบหมายงาน '{data.get('title', old_card['title'])}' ให้กับคุณ",
                link="#view=kanban"
            )
            batch_send_push_notification([a], "มอบหมายงาน / Kanban", f"{current_user}: {data.get('title', old_card['title'])}", url="#view=kanban")
            # --- Send to LINE ---
            threading.Thread(target=send_line_push_notification, args=(
                a, 
                "อัปเดตงาน (Kanban)", 
                f"{current_user} ได้เพิ่มคุณในงาน '{data.get('title', old_card['title'])}'"
            ), kwargs={
                "fields": {
                    "งาน": data.get('title', old_card['title']),
                    "แจ้งเตือน": f"{current_user} มอบหมายงานให้คุณ",
                    "ดูที่หน้าเว็บ": "เมนู Kanban"
                }
            }).start()
        except Exception as e:
            print(f"Error sending kanban update notification to {a}: {e}")
            
    return jsonify({"ok": True})

@app.route("/api/kanban/cards/<int:card_id>", methods=["DELETE"])
@login_required
def kanban_delete_card(card_id):
    ok = database.kanban_delete_card(card_id)
    return jsonify({"ok": ok})

@app.route("/api/kanban/cards/<int:card_id>/move", methods=["POST"])
@login_required
def kanban_move_card(card_id):
    data = request.json or {}
    new_column_id = data.get("column_id")
    new_position = data.get("position", 0)
    if new_column_id is None:
        return jsonify({"ok": False, "error": "Missing column_id"}), 400
    ok = database.kanban_move_card(card_id, int(new_column_id), int(new_position))
    
    # Also update 'is_done' if passed (for automatic status on move to Done column)
    is_done = data.get("is_done")
    if is_done is not None:
        database.kanban_update_card(card_id, is_done=int(is_done))
        
    # Broadcast board update via socketio
    socketio.emit("kanban_update", {}, room="kanban_board")
    return jsonify({"ok": ok})

@app.route("/api/users/list", methods=["GET"])
@login_required
def api_users_list():
    users = database.admin_get_all_users()
    # Mask notes and potentially other sensitive admin data for a general user list
    clean_users = []
    for u in users:
        clean_users.append({
            "username": u["username"],
            "display_name": u["display_name"],
            "avatar_url": u["avatar_url"],
            "department": u.get("department", "General")
        })
    return jsonify({"ok": True, "users": clean_users})

# ─────────────────────── Wiki API ─────────────────────────

@app.route("/api/wiki/pages", methods=["GET"])
@login_required
def wiki_list_pages():
    q = request.args.get("q", "").strip()
    if q:
        pages = database.wiki_search(q)
    else:
        pages = database.wiki_get_all_pages()
    return jsonify({"ok": True, "pages": pages})

@app.route("/api/wiki/pages", methods=["POST"])
@login_required
def wiki_create_page():
    data = request.json or {}
    title = data.get("title", "").strip()
    content = data.get("content", "")
    if not title:
        return jsonify({"ok": False, "error": "กรุณาระบุหัวข้อบทความ"}), 400
    page_id, slug = database.wiki_create_page(
        title=title,
        content=content,
        author=session.get("user", "Anonymous"),
        category_id=data.get("category_id")
    )
    return jsonify({"ok": True, "id": page_id, "slug": slug})

@app.route("/api/wiki/pages/<int:page_id>", methods=["GET"])
@login_required
def wiki_get_page(page_id):
    page = database.wiki_get_page(page_id=page_id)
    if not page:
        return jsonify({"ok": False, "error": "ไม่พบบทความ"}), 404
    return jsonify({"ok": True, "page": page})

@app.route("/api/wiki/pages/<int:page_id>", methods=["PUT"])
@login_required
def wiki_update_page(page_id):
    data = request.json or {}
    database.wiki_update_page(page_id, title=data.get("title"), content=data.get("content"), category_id=data.get("category_id"))
    return jsonify({"ok": True})

@app.route("/api/wiki/pages/<int:page_id>", methods=["DELETE"])
@login_required
def wiki_delete_page(page_id):
    ok = database.wiki_delete_page(page_id)
    return jsonify({"ok": ok})

try:
    from docx import Document
except ImportError:
    Document = None

@app.route("/api/wiki/pages/<int:page_id>/export", methods=["GET"])
@login_required
def export_wiki_page(page_id):
    export_format = request.args.get("format", "txt").lower()
    page = database.wiki_get_page(page_id=page_id)
    if not page:
        return jsonify({"ok": False, "error": "ไม่พบบทความ"}), 404

    title = page.get("title", "wiki_export")
    content = page.get("content", "")
    author = page.get("author", "Unknown")
    updated_at = page.get("updated_at", "")

    # Clean filename
    safe_title = "".join([c for c in title if c.isalnum() or c in [' ', '-', '_']]).strip().replace(" ", "_")
    if not safe_title:
        safe_title = f"wiki_page_{page_id}"

    downloads_dir = Path("static/downloads")
    downloads_dir.mkdir(parents=True, exist_ok=True)
    
    if export_format == "txt":
        file_path = downloads_dir / f"{safe_title}.txt"
        file_content = f"Title: {title}\nAuthor: {author}\nLast Updated: {updated_at}\n\n{content}"
        file_path.write_text(file_content, encoding="utf-8")
        return send_file(str(file_path), as_attachment=True, download_name=f"{safe_title}.txt")

    elif export_format == "pdf":
        file_path = downloads_dir / f"{safe_title}.pdf"
        pdf = FPDF()
        pdf.add_page()
        
        font_paths = [
            r"C:\Windows\Fonts\leelawad.ttf", r"C:\Windows\Fonts\thsarabunnew.ttf",
            r"C:\Windows\Fonts\tahoma.ttf", r"C:\Windows\Fonts\arial.ttf"
        ] if os.name == 'nt' else [
            "/usr/share/fonts/truetype/thai/THSarabunNew.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"
        ]
        
        font_path = next((f for f in font_paths if os.path.exists(f)), None)
            
        if font_path:
            pdf.add_font("THFont", "", font_path)
            pdf.set_font("THFont", size=16)
        else:
            pdf.set_font("Arial", size=16)

        pdf.cell(0, 10, txt=title, ln=1, align='C')
        pdf.set_font(pdf.font_family, '', 10)
        pdf.cell(0, 10, txt=f"Author: {author} | Updated: {updated_at}", ln=1, align='C')
        pdf.ln(5)
        pdf.set_font(pdf.font_family, '', 12)
        
        pdf.multi_cell(0, 7, txt=content)
        pdf.output(str(file_path))
        return send_file(str(file_path), as_attachment=True, download_name=f"{safe_title}.pdf")

    elif export_format == "docx":
        if Document is None:
            return jsonify({"ok": False, "error": "python-docx package is not installed."}), 400
        
        file_path = downloads_dir / f"{safe_title}.docx"
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(f"Author: {author} | Updated: {updated_at}")
        doc.add_paragraph(content)
        doc.save(str(file_path))
        
        return send_file(str(file_path), as_attachment=True, download_name=f"{safe_title}.docx")

    return jsonify({"ok": False, "error": "Invalid format"}), 400


if __name__ == "__main__":

    import os
    port = int(os.environ.get("PORT", 5005))
    
    # Optional sync on startup
    if os.environ.get("SKIP_SYNC") != "1":
        print("🚀 Syncing and re-indexing Knowledge Base (Universal)...", flush=True)
        try:
            rag_engine.fix_categories()
            rag_engine.sync_uploads()
        except Exception as e:
            print(f"⚠️ Startup background update failed: {e}", flush=True)
    else:
        print("⏭️  Skipping startup sync as requested.")
        
    print(f"🚀 Starting Org Chatbot at http://127.0.0.1:{port}")
    
    # Critical Diagnostic: Print all routes
    print("🗺️  Registered Routes:")
    for rule in app.url_map.iter_rules():
        print(f"   {rule.endpoint:20} -> {rule.rule}")
        
    # Run SocketIO server (disable reloader on Windows to prevent port conflicts)
    socketio.run(app, debug=False, port=port, host="0.0.0.0", use_reloader=False, allow_unsafe_werkzeug=True)
