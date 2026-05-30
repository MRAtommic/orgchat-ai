import os
import json
from pathlib import Path
from fpdf import FPDF

try:
    from docx import Document
except ImportError:
    Document = None

def export_page_to_file(page_id: int, title: str, content: str, author: str, updated_at: str, export_format: str) -> tuple[str | None, str | None, str | None]:
    """
    Formats and exports a wiki page to PDF, Docx, or TXT.
    Returns (file_path, download_name, error_message).
    """
    export_format = export_format.lower()
    safe_title = "".join([c for c in title if c.isalnum() or c in [' ', '-', '_']]).strip().replace(" ", "_")
    if not safe_title:
        safe_title = f"wiki_page_{page_id}"

    downloads_dir = Path("static/downloads")
    downloads_dir.mkdir(parents=True, exist_ok=True)

    if export_format == "txt":
        file_path = downloads_dir / f"{safe_title}.txt"
        file_content = f"Title: {title}\nAuthor: {author}\nLast Updated: {updated_at}\n\n{content}"
        file_path.write_text(file_content, encoding="utf-8")
        return str(file_path), f"{safe_title}.txt", None

    elif export_format == "pdf":
        file_path = downloads_dir / f"{safe_title}.pdf"
        # fpdf2 supports unicode and complex scripts better
        pdf = FPDF()
        pdf.add_page()
        
        # Common Thai fonts on Windows and Linux
        font_paths = [
            r"C:\Windows\Fonts\thsarabunnew.ttf", r"C:\Windows\Fonts\leelawad.ttf",
            r"C:\Windows\Fonts\tahoma.ttf", r"C:\Windows\Fonts\arial.ttf"
        ] if os.name == 'nt' else [
            "/usr/share/fonts/truetype/thai/THSarabunNew.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"
        ]
        
        font_path = next((f for f in font_paths if os.path.exists(f)), None)
            
        if font_path:
            # fpdf2: add_font automatically handles unicode
            pdf.add_font("THFont", "", font_path)
            pdf.set_font("THFont", size=16)
        else:
            pdf.set_font("helvetica", size=16)

        # Title
        pdf.cell(0, 10, text=title, new_x="LMARGIN", new_y="NEXT", align='C')
        
        # Meta info
        if font_path: pdf.set_font("THFont", size=10)
        else: pdf.set_font("helvetica", size=10)
        pdf.cell(0, 10, text=f"Author: {author} | Updated: {updated_at}", new_x="LMARGIN", new_y="NEXT", align='C')
        pdf.ln(5)
        
        # Content
        if font_path: pdf.set_font("THFont", size=12)
        else: pdf.set_font("helvetica", size=12)
        
        # Replace complex newlines to avoid issues
        clean_content = content.replace("\r\n", "\n")
        pdf.multi_cell(0, 7, text=clean_content)
        
        pdf.output(str(file_path))
        return str(file_path), f"{safe_title}.pdf", None

    elif export_format == "docx":
        if Document is None:
            return None, None, "python-docx package is not installed."
        
        file_path = downloads_dir / f"{safe_title}.docx"
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(f"Author: {author} | Updated: {updated_at}")
        doc.add_paragraph(content)
        doc.save(str(file_path))
        return str(file_path), f"{safe_title}.docx", None

    return None, None, "Invalid export format"
